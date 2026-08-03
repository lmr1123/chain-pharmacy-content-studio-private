import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


MODULE_PATH = Path(__file__).with_name("import_courseware_content.py")
SPEC = importlib.util.spec_from_file_location("courseware_importer", MODULE_PATH)
IMPORTER = importlib.util.module_from_spec(SPEC)
assert SPEC.loader
sys.modules[SPEC.name] = IMPORTER
SPEC.loader.exec_module(IMPORTER)


class CoursewareImporterTest(unittest.TestCase):
    def parse(self, text):
        return IMPORTER.parse_blocks(
            IMPORTER.text_blocks(text),
            "sample.txt",
            "text",
        )

    def test_complete_headed_text_and_overflow(self):
        manifest = self.parse(
            """商品名称：测试商品
规格：100ml
一、商品介绍
主要成分：示例成分
功能主治：示例原文
二、核心卖点
1、卖点一：原文一
2、卖点二：原文二
3、卖点三：原文三
4、卖点四：原文四
三、适宜人群
1、示例人群一
2、示例人群二
四、联合用药
应用场景：示例场景
联合用药：测试商品 + 联合商品
销售话术：示例话术
五、注意事项
1、过敏者慎用。"""
        )
        self.assertEqual(manifest["product"]["display_name"], "测试商品")
        self.assertEqual(len(manifest["selling_points"]), 4)
        self.assertEqual(len(manifest["combinations"]), 1)
        self.assertEqual(manifest["page_rules"]["page01"]["page_count"], 2)
        self.assertEqual(manifest["page_rules"]["page02"]["page_count"], 1)
        self.assertEqual(manifest["combinations"][0]["talk_track"], "示例话术")
        self.assertGreater(len(manifest["source_blocks"]), 10)

    def test_missing_sections_stay_blank(self):
        manifest = self.parse(
            """商品名称：缺项商品
商品介绍
主要成分：示例成分
注意事项
儿童应在医师指导下使用。
其他内容
无法归类原文。"""
        )
        self.assertEqual(manifest["selling_points"], [])
        self.assertEqual(manifest["audiences"], [])
        self.assertEqual(manifest["combinations"], [])
        self.assertIn("selling_points", manifest["blank_fields"])
        self.assertIn("audiences", manifest["blank_fields"])
        self.assertIn("combinations", manifest["blank_fields"])
        self.assertEqual(manifest["unmapped_content"][0]["text"], "无法归类原文。")

    def test_unheaded_content_is_not_invented(self):
        manifest = self.parse(
            """自然段商品
主要成分：示例成分
这是没有可靠栏目线索的一段原文。
对本品过敏者禁用。"""
        )
        self.assertEqual(manifest["product"]["display_name"], "自然段商品")
        self.assertEqual(manifest["introduction"][0]["field_name"], "主要成分")
        self.assertEqual(
            manifest["unmapped_content"][0]["text"],
            "这是没有可靠栏目线索的一段原文。",
        )
        self.assertEqual(
            manifest["precautions"][0]["content"],
            "对本品过敏者禁用。",
        )

    def test_docx_key_value_and_combination_tables(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "table.docx"
            document = Document()
            table = document.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "商品名称"
            table.cell(0, 1).text = "Word表格商品"
            table.cell(1, 0).text = "规格"
            table.cell(1, 1).text = "10片"
            table.cell(2, 0).text = "主要成分"
            table.cell(2, 1).text = "示例成分"
            combo = document.add_table(rows=2, cols=4)
            headers = ["应用场景", "联合用药", "产品图片展示", "销售话术"]
            values = ["示例场景", "商品A + 商品B", "", "示例话术"]
            for index, value in enumerate(headers):
                combo.cell(0, index).text = value
            for index, value in enumerate(values):
                combo.cell(1, index).text = value
            document.save(path)

            blocks, source_type = IMPORTER.read_source(path)
            manifest = IMPORTER.parse_blocks(blocks, path.name, source_type)
            self.assertEqual(manifest["product"]["display_name"], "Word表格商品")
            self.assertEqual(manifest["product"]["specification"], "10片")
            self.assertEqual(manifest["introduction"][0]["content"], "示例成分")
            self.assertEqual(manifest["combinations"][0]["scenario"], "示例场景")
            self.assertEqual(manifest["combinations"][0]["product_image"], "")

    def test_benchmark_table_is_preserved(self):
        manifest = IMPORTER.parse_blocks(
            [
                IMPORTER.Block(
                    kind="table",
                    headers=["对比维度", "本品", "竞品"],
                    rows=[["零售价", "10元", "12元"]],
                )
            ],
            "benchmark.docx",
            "docx",
        )
        self.assertEqual(manifest["benchmarks"][0]["headers"][0], "对比维度")
        self.assertEqual(manifest["benchmarks"][0]["rows"][0][1], "10元")

    def test_image_request_builds_guarded_prompt(self):
        manifest = IMPORTER.parse_blocks(
            [
                IMPORTER.Block(
                    kind="table",
                    headers=["使用位置", "图片主题（只写一句话）", "补充要求"],
                    rows=[
                        [
                            "04 注意事项",
                            "药师提醒成年患者清淡饮食",
                            "人物动作自然",
                        ]
                    ],
                )
            ],
            "image-request.docx",
            "docx",
        )
        request = manifest["image_requests"][0]
        self.assertEqual(request["mode"], "auto-or-prompt")
        self.assertIn("药师提醒成年患者清淡饮食", request["prompt"])
        self.assertIn("不生成任何文字", request["prompt"])
        self.assertTrue(manifest["image_generation_policy"]["medical_review_required"])

    def test_real_product_asset_is_not_generated(self):
        manifest = IMPORTER.parse_blocks(
            [
                IMPORTER.Block(
                    kind="table",
                    headers=["使用位置", "图片主题"],
                    rows=[["商品主图", "生成某品牌真实药盒"]],
                )
            ],
            "product-image.docx",
            "docx",
        )
        request = manifest["image_requests"][0]
        self.assertEqual(request["mode"], "authorized-asset-required")
        self.assertEqual(request["prompt"], "")
        self.assertTrue(request["requires_authorized_asset"])

    def test_standardized_five_slide_profile_selects_registered_template(self):
        manifest = IMPORTER.parse_blocks(
            [
                IMPORTER.Block(
                    kind="table",
                    headers=["填写项目", "业务填写内容"],
                    rows=[
                        ["课件类型（请勿修改）", "简版商品培训课件（5页）"],
                        ["课程标题", "测试商品培训"],
                        ["商品名称", "测试商品"],
                    ],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["商品字段", "审核原文"],
                    rows=[["主要成分", "测试成分"]],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["卖点名称", "支撑内容"],
                    rows=[["测试卖点", "测试依据"]],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["适宜人群"],
                    rows=[["测试人群"]],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["应用场景", "联合用药", "销售话术"],
                    rows=[["测试场景", "商品A + 商品B", "测试话术"]],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["对比维度", "本品", "竞品"],
                    rows=[["规格", "10ml", "20ml"]],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["注意事项"],
                    rows=[["测试注意事项"]],
                ),
            ],
            "five-slide.docx",
            "docx",
        )
        self.assertEqual(
            manifest["template_id"],
            "template.product-courseware-dashenlin-green-v1",
        )
        self.assertEqual(manifest["courseware_profile"], "product-brief-5")
        self.assertEqual(manifest["product"]["display_name"], "测试商品")
        self.assertEqual(len(manifest["combinations"]), 1)
        self.assertEqual(manifest["unmapped_content"], [])

    def test_standardized_disease_product_profile_keeps_scenario_structure(self):
        manifest = IMPORTER.parse_blocks(
            [
                IMPORTER.Block(
                    kind="table",
                    headers=["填写项目", "业务填写内容"],
                    rows=[
                        [
                            "课件类型（请勿修改）",
                            "疾病—商品—场景培训课件（18页）",
                        ],
                        ["课程标题", "测试课程"],
                        ["疾病主题", "测试证型"],
                        ["主推商品", "测试商品"],
                        ["培训对象", "门店员工"],
                        ["一句话导语", "测试导语"],
                    ],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["填写项目", "业务填写内容"],
                    rows=[
                        ["疾病定义", "测试定义"],
                        ["通俗说明", "测试说明"],
                        ["主要辨证要点", "测试要点"],
                    ],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=[
                        "场景名称",
                        "辨证沟通",
                        "核心用药",
                        "关联服务",
                        "服务要点",
                    ],
                    rows=[
                        [
                            "测试场景",
                            "测试沟通",
                            "测试用药",
                            "测试服务",
                            "测试要点",
                        ]
                    ],
                ),
            ],
            "disease-product.docx",
            "docx",
        )
        self.assertEqual(
            manifest["template_id"],
            "template.dashenlin-disease-product-scenario-v1",
        )
        self.assertEqual(
            manifest["courseware_profile"], "disease-product-scenario-18"
        )
        self.assertEqual(manifest["course"]["disease_theme"], "测试证型")
        self.assertEqual(manifest["product"]["display_name"], "测试商品")
        self.assertEqual(
            manifest["content_sections"]["scenario_solutions"][0]["场景名称"],
            "测试场景",
        )
        self.assertEqual(manifest["page_rules"]["scenarios"]["page_count"], 1)

    def test_unreplaced_business_placeholder_is_treated_as_blank(self):
        manifest = IMPORTER.parse_blocks(
            [
                IMPORTER.Block(
                    kind="table",
                    headers=["填写项目", "业务填写内容"],
                    rows=[
                        ["课件类型（请勿修改）", "简版商品培训课件（5页）"],
                        ["商品名称", "测试商品"],
                    ],
                ),
                IMPORTER.Block(
                    kind="table",
                    headers=["卖点名称", "支撑内容"],
                    rows=[
                        [
                            "测试卖点",
                            "【业务填写】粘贴公司已审核终稿；无内容可留空。",
                        ]
                    ],
                ),
            ],
            "placeholder.docx",
            "docx",
        )
        self.assertEqual(manifest["selling_points"][0]["content"], "")


if __name__ == "__main__":
    unittest.main()
