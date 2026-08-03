from __future__ import annotations

import base64
import tempfile
import unittest
from pathlib import Path

from docx import Document

from build_universal_courseware_business_word import build
from import_universal_courseware_content import parse_docx


ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "YAAAAAYAAjCB0C8AAAAASUVORK5CYII="
)


class UniversalCoursewareInputTest(unittest.TestCase):
    def test_blank_template_does_not_create_fake_content(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = root / "template.docx"
            build().save(path)
            manifest = parse_docx(path, root / "assets")
            self.assertEqual(manifest["course"]["theme"], "")
            self.assertEqual(manifest["sections"], [])
            self.assertEqual(manifest["content_metrics"]["image_count"], 0)

    def test_variable_item_counts_produce_layout_candidates_without_empty_cards(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = root / "filled.docx"
            document = Document()
            document.add_paragraph("课件主题：春季过敏健康培训")
            document.add_paragraph("培训对象：门店员工")
            document.add_heading("典型表现", level=1)
            document.add_paragraph("鼻痒、连续喷嚏。")
            document.add_paragraph("清水样鼻涕。")
            document.add_heading("日常建议", level=1)
            for item in ("减少暴露", "及时清洁", "规律作息", "适度运动", "必要时就医"):
                document.add_paragraph(item)
            document.save(path)

            manifest = parse_docx(path, root / "assets")
            self.assertEqual(manifest["content_metrics"]["section_count"], 2)
            self.assertIn("two_card", manifest["sections"][0]["layout_candidates"])
            self.assertIn(
                "five_card_3_plus_2",
                manifest["sections"][1]["layout_candidates"],
            )
            self.assertEqual(
                manifest["planning_policy"]["empty_cards"],
                "forbidden",
            )
            self.assertFalse(manifest["planning_policy"]["fixed_item_count"])

    def test_pasted_image_stays_with_its_section(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            image_path = root / "symptom.png"
            image_path.write_bytes(ONE_PIXEL_PNG)
            path = root / "image.docx"
            document = Document()
            document.add_paragraph("课件主题：图片测试")
            document.add_heading("症状说明", level=1)
            document.add_paragraph("审核正文")
            paragraph = document.add_paragraph()
            paragraph.add_run().add_picture(str(image_path))
            document.add_paragraph("图片说明／来源：公司授权症状图")
            document.save(path)

            manifest = parse_docx(path, root / "assets")
            image = manifest["sections"][0]["images"][0]
            self.assertEqual(manifest["content_metrics"]["image_count"], 1)
            self.assertEqual(image["caption_or_source"], "公司授权症状图")
            self.assertTrue(Path(image["asset_path"]).exists())

    def test_shared_caption_applies_to_two_images_in_same_paragraph(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            image_path = root / "symptom.png"
            image_path.write_bytes(ONE_PIXEL_PNG)
            path = root / "two-images.docx"
            document = Document()
            document.add_paragraph("课件主题：双图测试")
            document.add_heading("典型表现", level=1)
            document.add_paragraph("发热。")
            document.add_paragraph("喉咙肿痛。")
            paragraph = document.add_paragraph()
            paragraph.add_run().add_picture(str(image_path))
            paragraph.add_run().add_picture(str(image_path))
            document.add_paragraph("图片说明／来源：公司授权双图")
            document.save(path)

            manifest = parse_docx(path, root / "assets")
            images = manifest["sections"][0]["images"]
            self.assertEqual(len(images), 2)
            self.assertEqual(
                [image["caption_or_source"] for image in images],
                ["公司授权双图", "公司授权双图"],
            )


if __name__ == "__main__":
    unittest.main()
