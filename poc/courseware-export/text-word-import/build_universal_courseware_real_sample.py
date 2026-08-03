#!/usr/bin/env python3
"""Build a real filled example for the universal courseware Word input."""

from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from build_universal_courseware_business_word import (
    BLACK,
    BLUE,
    MUTED,
    REPO_DIR,
    create_document,
    set_run_font,
)


OUTPUT = (
    REPO_DIR
    / "outputs/courseware-natural-import/风热证培训课件内容与素材提交_真实已填样本.docx"
)
FEVER_IMAGE = (
    REPO_DIR / "assets/component-library/symptoms/fever/master/fever-v1.png"
)
THROAT_IMAGE = (
    REPO_DIR
    / "assets/component-library/symptoms/sore-throat/master/sore-throat-v1.png"
)


def add_meta(document, label: str, value: str):
    paragraph = document.add_paragraph()
    set_run_font(paragraph.add_run(f"{label}："), size=11, bold=True)
    set_run_font(paragraph.add_run(value), size=11)


def add_section(document, title: str, paragraphs: list[str]):
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(title)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = False
        set_run_font(paragraph.add_run(text), size=11)


def add_image_pair(
    document,
    left_path: Path,
    left_alt: str,
    right_path: Path,
    right_alt: str,
    caption: str,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    for path, alt in ((left_path, left_alt), (right_path, right_alt)):
        run = paragraph.add_run()
        picture = run.add_picture(str(path), width=Inches(2.05))
        picture._inline.docPr.set("descr", alt)
        picture._inline.docPr.set("title", alt)
        paragraph.add_run("  ")

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.keep_with_next = False
    set_run_font(
        caption_paragraph.add_run(f"图片说明／来源：{caption}"),
        size=9,
        color=MUTED,
        italic=True,
    )


def build():
    document = create_document()

    title = document.add_paragraph()
    title.paragraph_format.space_after = Inches(0.04)
    set_run_font(
        title.add_run("培训课件内容与素材提交"),
        size=24,
        color=BLACK,
        bold=True,
    )

    subtitle = document.add_paragraph(style="Universal Ignore")
    set_run_font(
        subtitle.add_run("真实已填样本｜风热证门店培训"),
        size=11,
        color=BLUE,
        bold=True,
    )

    notice = document.add_paragraph(style="Universal Ignore")
    set_run_font(
        notice.add_run(
            "以下内容取自项目现有内部样本，仅用于展示业务如何整理内容和粘贴素材；"
            "新项目投产前仍需完成医学、药事、合规和素材授权审核。"
        ),
        size=9.5,
        color=MUTED,
    )

    add_meta(document, "课件主题", "中医基础知识——风热证")
    add_meta(document, "培训对象（可不填）", "连锁药店一线员工")
    add_meta(
        document,
        "培训目标（可不填）",
        "理解风热证的形成与典型表现，掌握基本调理思路和日常注意事项。",
    )

    add_section(
        document,
        "什么是风热证",
        [
            "简单来说，风热证就是风邪和热邪一起侵入身体，导致体表不适、肺气不顺，因此可能出现发热、咽痛、口渴等偏“热”的表现。"
        ],
    )

    add_section(
        document,
        "典型表现",
        [
            "发热、口渴：身体发热，嘴巴容易觉得干渴，也可能伴有烦躁。",
            "喉咙肿痛：咽喉红肿疼痛，还可能伴有咳嗽、痰黄或黄稠鼻涕。",
        ],
    )
    add_image_pair(
        document,
        FEVER_IMAGE,
        "发热症状插画",
        THROAT_IMAGE,
        "喉咙肿痛症状插画",
        "从左至右：发热、喉咙肿痛；项目内原创素材 symptom.fever、symptom.sore-throat",
    )

    add_section(
        document,
        "调理思路",
        [
            "核心思路是疏风清热。",
            "日常可选用桑叶、菊花、薄荷等内容进行知识讲解；涉及用量或饮用方法时，应以公司审核终稿为准。",
            "如需展示药品名称或包装，应同时提交公司审核文案和授权包装原图；未提供时不展示包装。",
        ],
    )

    add_section(
        document,
        "日常注意事项",
        [
            "房间适当开窗通风，保持空气流通。",
            "少量多次饮用温水，注意补充水分。",
            "饮食以清淡为主，少吃辛辣、油炸和燥热食物。",
            "暂时避免烟酒，并保证充足休息。",
        ],
    )

    add_section(
        document,
        "课程小结",
        [
            "风热证的学习重点可以概括为：病因看“风邪＋热邪”，表现看“发热、口渴、咽痛等偏热症状”，调理把握“疏风清热”，生活中注意通风、补水、清淡饮食和休息。"
        ],
    )

    document.core_properties.title = "风热证培训课件内容与素材提交真实已填样本"
    document.core_properties.subject = "通用内容驱动课件 Word 的真实填写示例"
    document.core_properties.keywords = "风热证, 培训课件, 真实样本, 内容驱动, 图片素材"
    return document


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
