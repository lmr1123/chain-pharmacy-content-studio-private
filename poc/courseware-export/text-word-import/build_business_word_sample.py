#!/usr/bin/env python3
"""Build the business-facing Word sample for natural courseware import."""

from __future__ import annotations

import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_DIR = Path("/Users/liminrong/Projects/chain-pharmacy-content-studio")
OUTPUT_PATH = (
    REPO_DIR
    / "outputs/courseware-natural-import/商品培训课件_业务填写样本.docx"
)
EXAMPLE_IMAGE_PATH = (
    REPO_DIR
    / "assets/component-library/advice-icons/light-diet-dashenlin/"
    "candidates/light-diet-dashenlin-v1.png"
)
SKILL_SCRIPTS = Path(
    "/Users/liminrong/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.727.11326/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402
sys.path.insert(0, str(Path(__file__).resolve().parent))
from image_prompt_protocol import copyable_prompt_template  # noqa: E402


GREEN = "009900"
DEEP_GREEN = "006B3C"
PALE_GREEN = "E7F3E2"
LIGHT_GREEN = "F3F8F1"
LIGHT_GRAY = "F4F6F5"
MID_GRAY = "66736B"
TEXT = "1B2E24"
WHITE = "FFFFFF"
LINE = "C7D8C2"
FONT_EAST_ASIA = "Source Han Sans SC"
FONT_ASCII = "Source Han Sans SC"


def set_run_font(
    run,
    *,
    size: float = 11,
    color: str = TEXT,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = FONT_ASCII
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_ASCII)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_ASCII)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_left_border(paragraph, color: str = GREEN, size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, value, fld_end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MID_GRAY)


def restart_numbering(document: Document, paragraphs: list) -> None:
    numbering = document.part.numbering_part.element
    base_num_id = document.styles["List Number"].element.pPr.numPr.numId.val
    base_num = next(
        num
        for num in numbering.findall(qn("w:num"))
        if int(num.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = (
        max(int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num")))
        + 1
    )
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    new_num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    new_num.append(override)
    numbering.append(new_num)

    for paragraph in paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        ilvl = num_pr.get_or_add_ilvl()
        ilvl.set(qn("w:val"), "0")
        num_id = num_pr.get_or_add_numId()
        num_id.set(qn("w:val"), str(new_num_id))


def add_hint(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="填写提示")
    paragraph.paragraph_format.left_indent = Inches(0.1)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(paragraph, PALE_GREEN)
    set_left_border(paragraph)
    run = paragraph.add_run(f"填写提示｜{text}")
    set_run_font(run, size=9.5, color=DEEP_GREEN)


def add_ignore_paragraph(
    document: Document,
    text: str,
    *,
    size: float = 11,
    color: str = MID_GRAY,
    bold: bool = False,
    after: float = 4,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    paragraph = document.add_paragraph(style="Courseware Ignore")
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_data_paragraph(
    document: Document,
    label: str,
    value: str,
    *,
    numbered: bool = False,
) -> None:
    paragraph = document.add_paragraph(style="List Number" if numbered else "Normal")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if numbered:
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run)


def format_table_text(table, *, header: bool = False, body_size: float = 10) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if header and row_index == 0:
                set_cell_fill(cell, GREEN)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9.5 if header and row_index == 0 else body_size,
                        color=WHITE if header and row_index == 0 else TEXT,
                        bold=header and row_index == 0,
                    )


def build_document() -> Document:
    document = Document()
    properties = document.core_properties
    properties.title = "商品培训课件业务填写 Word 样本"
    properties.subject = "文本与 Word 自动整理导入"
    properties.author = "大参林医药集团"
    properties.last_modified_by = ""
    properties.comments = "一份 Word 对应一个商品；示例内容不可直接作为生产课件文案。"
    properties.created = datetime(2026, 7, 30, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 7, 30, tzinfo=timezone.utc)
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_ASCII
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 12, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = FONT_ASCII
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(DEEP_GREEN)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    ignore_style = styles.add_style("Courseware Ignore", WD_STYLE_TYPE.PARAGRAPH)
    ignore_style.base_style = normal
    hint_style = styles.add_style("填写提示", WD_STYLE_TYPE.PARAGRAPH)
    hint_style.base_style = normal

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    left = header.add_run("大参林｜商品培训课件内容导入")
    set_run_font(left, size=9, color=DEEP_GREEN, bold=True)
    header.add_run("\t")
    right = header.add_run("业务填写样本")
    set_run_font(right, size=9, color=MID_GRAY)
    tabs = header.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5))

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    add_ignore_paragraph(
        document,
        "商品培训课件内容导入",
        size=11,
        color=GREEN,
        bold=True,
        after=4,
    )
    title = document.add_paragraph(style="Courseware Ignore")
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("业务填写 Word 样本")
    set_run_font(title_run, size=24, color=DEEP_GREEN, bold=True)
    add_ignore_paragraph(
        document,
        "一份 Word 对应一个商品｜系统自动整理栏目｜缺失内容可留空",
        size=12,
        color=MID_GRAY,
        after=10,
    )

    add_hint(
        document,
        "绿色提示和本页使用说明不会进入课件内容。业务只需要替换示例文字；没有的内容直接留空，不需要补齐所有栏目。",
    )

    add_ignore_paragraph(document, "使用方法", size=13, color=DEEP_GREEN, bold=True, after=4)
    for text in (
        "保留下面的栏目标题，直接替换示例内容。",
        "商品介绍、卖点、人群和联合用药条数可不同；缺失内容直接留空。",
        "功效、用法、联合用药和销售话术只粘贴内部已审核原文。",
    ):
        paragraph = document.add_paragraph(style="Courseware Ignore")
        paragraph.style = styles["Courseware Ignore"]
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(f"• {text}")
        set_run_font(run, size=10, color=TEXT)

    add_ignore_paragraph(
        document,
        "商品基本信息",
        size=13,
        color=DEEP_GREEN,
        bold=True,
        after=4,
    )
    add_hint(
        document,
        "右侧单元格填写内容；商品主图可暂时留空，后续在 PPTX 中粘贴公司授权原图。真实包装不得用 AI 仿造。",
    )
    metadata = [
        ("商品名称", "示例商品A（演示用）"),
        ("品牌名称", "大参林"),
        ("商品编码", "DEMO-WORD-001"),
        ("主推", "A"),
        ("规格", "示例规格"),
        ("零售价", "示例价格"),
        ("一句话卖点", "示例：请替换为内部已审核的一句话卖点"),
        ("商品主图", ""),
    ]
    table = document.add_table(rows=len(metadata), cols=2)
    for row_index, (label, value) in enumerate(metadata):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
        set_cell_fill(table.cell(row_index, 0), LIGHT_GREEN)
        for run in table.cell(row_index, 0).paragraphs[0].runs:
            set_run_font(run, size=9.5, color=DEEP_GREEN, bold=True)
    apply_table_geometry(table, [2700, 6660])
    format_table_text(table, body_size=9.5)

    document.add_page_break()

    document.add_heading("一、商品介绍", level=1)
    add_hint(document, "可填写主要成分、功能主治、用法用量等；没有的字段直接删除或留空。")
    add_data_paragraph(document, "主要成分", "示例成分，请替换为内部审核原文")
    add_data_paragraph(document, "功能主治", "示例功效描述，请替换为内部审核原文")
    add_data_paragraph(document, "用法用量", "示例用法用量，请替换为内部审核原文")

    document.add_heading("二、核心卖点", level=1)
    add_hint(document, "一条卖点使用一个编号段落；条数不限，超过单页容量时系统复制原 01 页续页。")
    selling_paragraphs = []
    for label, value in (
        ("示例卖点一", "请填写卖点说明或支撑信息。"),
        ("示例卖点二", "不同商品的卖点数量可以不同。"),
        ("示例卖点三", "请保留内部审核后的原始表述。"),
        ("示例卖点四", "第四条用于演示内容超量时自动生成 01 续页。"),
    ):
        add_data_paragraph(document, label, value, numbered=True)
        selling_paragraphs.append(document.paragraphs[-1])
    restart_numbering(document, selling_paragraphs)

    document.add_heading("三、适宜人群", level=1)
    add_hint(document, "一类人群使用一个编号段落；如果没有明确资料，本节可以留空。")
    audience_paragraphs = []
    for audience in (
        "示例适宜人群一",
        "示例适宜人群二",
        "示例适宜人群三",
    ):
        paragraph = document.add_paragraph(audience, style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        for run in paragraph.runs:
            set_run_font(run)
        audience_paragraphs.append(paragraph)
    restart_numbering(document, audience_paragraphs)

    document.add_page_break()

    document.add_heading("四、联合用药", level=1)
    add_hint(
        document,
        "每一行代表一个应用场景。产品图片展示可以留空，生成 PPTX 后再放入商品包装图；话术只填写内部审核稿。",
    )
    combo_headers = ["应用场景", "联合用药", "产品图片展示", "销售话术"]
    combo_rows = [
        [
            "示例应用场景一",
            "示例商品A + 联合商品一",
            "",
            "示例话术，请替换为内部审核原文。",
        ],
        [
            "示例应用场景二",
            "示例商品A + 联合商品二",
            "",
            "",
        ],
    ]
    combo = document.add_table(rows=1 + len(combo_rows), cols=4)
    for col_index, value in enumerate(combo_headers):
        combo.cell(0, col_index).text = value
    for row_index, values in enumerate(combo_rows, start=1):
        for col_index, value in enumerate(values):
            combo.cell(row_index, col_index).text = value
    set_repeat_table_header(combo.rows[0])
    apply_table_geometry(combo, [1800, 2300, 1900, 3360])
    format_table_text(combo, header=True, body_size=9.5)

    document.add_heading("五、品种对标", level=1)
    add_hint(document, "只有确实需要竞品对比时填写；没有资料可以整节留空。")
    benchmark_headers = ["对比维度", "本品", "竞品"]
    benchmark_rows = [
        ["功效主治", "示例本品内容", "示例竞品内容"],
        ["零售价", "示例价格", "示例价格"],
        ["卖点差异", "示例差异", "示例差异"],
    ]
    benchmark = document.add_table(rows=1 + len(benchmark_rows), cols=3)
    for col_index, value in enumerate(benchmark_headers):
        benchmark.cell(0, col_index).text = value
    for row_index, values in enumerate(benchmark_rows, start=1):
        for col_index, value in enumerate(values):
            benchmark.cell(row_index, col_index).text = value
    set_repeat_table_header(benchmark.rows[0])
    apply_table_geometry(benchmark, [1900, 3730, 3730])
    format_table_text(benchmark, header=True, body_size=10)

    document.add_page_break()

    document.add_heading("六、注意事项", level=1)
    add_hint(document, "一条注意事项使用一个编号段落；必须以说明书或内部审核稿为准。")
    precaution_paragraphs = []
    for content in (
        "示例注意事项一，请替换为审核原文。",
        "示例注意事项二，请替换为审核原文。",
        "示例注意事项三，请替换为审核原文。",
        "示例注意事项四，请替换为审核原文。",
    ):
        paragraph = document.add_paragraph(content, style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        for run in paragraph.runs:
            set_run_font(run)
        precaution_paragraphs.append(paragraph)
    restart_numbering(document, precaution_paragraphs)

    document.add_heading("七、其他内容", level=1)
    add_hint(
        document,
        "不知道应该放在哪个栏目时，可把原文放在这里。系统会将其列入“待确认内容”，不会擅自归类或丢弃。",
    )
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("示例：这是一段暂时无法判断栏目归属的原文。")
    set_run_font(run)

    add_ignore_paragraph(document, "提交前自查", size=15, color=DEEP_GREEN, bold=True, after=8)
    for item in (
        "一份 Word 只包含一个商品。",
        "商品名称已填写；其余没有资料的栏目允许留空。",
        "功效、用法、联合用药和销售话术均为内部审核原文。",
        "商品图、联合商品图使用公司授权原图；内容插画参考下一页单独生成。",
        "文件名建议使用“商品名称_版本日期.docx”。",
    ):
        add_ignore_paragraph(document, f"□ {item}", size=10.5, color=TEXT, after=4)

    add_hint(
        document,
        "本文件中的“示例”内容仅用于说明整理方式，不可直接作为生产课件文案。",
    )

    document.add_page_break()

    add_ignore_paragraph(
        document,
        "图片生成与提交",
        size=18,
        color=DEEP_GREEN,
        bold=True,
        after=6,
    )
    add_hint(
        document,
        "最简单的做法：填图片主题 → 复制提示词到任意生图系统（或由本项目自动生成）→ 人工调整确认 → 把最终 PNG 粘贴到 PPTX。正文不用重复填写图片说明。",
    )
    add_ignore_paragraph(
        document,
        "第一步｜填写图片主题",
        size=12,
        color=DEEP_GREEN,
        bold=True,
        after=3,
    )
    add_ignore_paragraph(
        document,
        "一句话写清“谁／什么，在做什么”。不要在这里填写药学结论；需要多张图片时继续增加行。",
        size=9,
        color=TEXT,
        after=4,
    )
    request_headers = [
        "使用位置",
        "图片主题（只写一句话）",
        "补充要求（可留空）",
        "案例示意图",
    ]
    request_rows = [
        [
            "04 注意事项",
            "药师提醒成年患者清淡饮食，餐桌上有蔬菜和清淡食物",
            "人物动作自然，画面亲和克制",
            "",
        ],
    ]
    requests = document.add_table(rows=1 + len(request_rows), cols=4)
    for col_index, value in enumerate(request_headers):
        requests.cell(0, col_index).text = value
    for row_index, values in enumerate(request_rows, start=1):
        for col_index, value in enumerate(values):
            requests.cell(row_index, col_index).text = value
    sample_cell = requests.cell(1, 3)
    sample_paragraph = sample_cell.paragraphs[0]
    sample_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sample_shape = sample_paragraph.add_run().add_picture(
        str(EXAMPLE_IMAGE_PATH),
        width=Inches(0.78),
    )
    sample_shape._inline.docPr.set(
        "descr",
        "候选风格示意：药师提醒成年患者清淡饮食，非医学证据图",
    )
    sample_caption = sample_cell.add_paragraph()
    sample_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sample_caption.paragraph_format.space_before = Pt(1)
    sample_caption.paragraph_format.space_after = Pt(0)
    caption_run = sample_caption.add_run("候选示意｜非证据图")
    set_run_font(caption_run, size=6.4, color=MID_GRAY)
    set_repeat_table_header(requests.rows[0])
    apply_table_geometry(requests, [1400, 3840, 2200, 1920])
    format_table_text(requests, header=True, body_size=8.5)

    add_ignore_paragraph(
        document,
        "第二步｜复制下面整段提示词",
        size=12,
        color=DEEP_GREEN,
        bold=True,
        after=3,
    )
    add_ignore_paragraph(
        document,
        "把【填写使用位置】【填写图片主题】替换成上表内容；没有补充要求就删除对应一行。",
        size=9,
        color=TEXT,
        after=3,
    )
    prompt_paragraph = document.add_paragraph(style="Courseware Ignore")
    prompt_paragraph.paragraph_format.left_indent = Inches(0.1)
    prompt_paragraph.paragraph_format.right_indent = Inches(0.1)
    prompt_paragraph.paragraph_format.space_before = Pt(0)
    prompt_paragraph.paragraph_format.space_after = Pt(3)
    prompt_paragraph.paragraph_format.line_spacing = 1.03
    set_paragraph_shading(prompt_paragraph, LIGHT_GRAY)
    set_left_border(prompt_paragraph, color=DEEP_GREEN, size="14")
    prompt_run = prompt_paragraph.add_run(copyable_prompt_template())
    set_run_font(prompt_run, size=7.8, color=TEXT)

    add_ignore_paragraph(
        document,
        "第三步｜确认后再粘贴到 PPTX",
        size=12,
        color=DEEP_GREEN,
        bold=True,
        after=3,
    )
    for item in (
        "主题与审核原文一致，不新增医学结论；AI 图仅作辅助，须经药师和视觉复核。",
        "图片中没有文字、药名、剂量、Logo、水印或仿造的真实包装。",
        "人物肢体、手指、医疗物件和医患动作自然，无明显生成错误。",
        "使用方形高清 PNG；建议不低于 1536×1536，四周保留安全区。",
    ):
        add_ignore_paragraph(document, f"□ {item}", size=8.5, color=TEXT, after=2)

    warning = document.add_paragraph(style="填写提示")
    warning.paragraph_format.space_before = Pt(2)
    warning.paragraph_format.space_after = Pt(0)
    warning.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(warning, "FFF2E8")
    set_left_border(warning, color="D6452D", size="18")
    warning_run = warning.add_run(
        "严谨要求｜商品包装、说明书、检测报告、处方和品牌证据必须使用公司授权原图，禁止 AI 仿造。"
    )
    set_run_font(warning_run, size=8.8, color="9C2B1E", bold=True)
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
