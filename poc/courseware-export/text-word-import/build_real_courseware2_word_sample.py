#!/usr/bin/env python3
"""Build a real-content business Word sample for courseware template 2."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_business_word_sample import (
    DEEP_GREEN,
    FONT_ASCII,
    FONT_EAST_ASIA,
    GREEN,
    LIGHT_GREEN,
    MID_GRAY,
    PALE_GREEN,
    TEXT,
    add_data_paragraph,
    add_hint,
    add_ignore_paragraph,
    add_page_number,
    apply_table_geometry,
    format_table_text,
    restart_numbering,
    set_cell_fill,
    set_repeat_table_header,
    set_run_font,
)


REPO_DIR = Path("/Users/liminrong/Projects/chain-pharmacy-content-studio")
OUTPUT_PATH = (
    REPO_DIR
    / "outputs/courseware-natural-import/"
    "穿心莲内酯滴丸_商品培训课件_业务真实内容样本.docx"
)
PACKSHOT_PATH = (
    REPO_DIR
    / "poc/courseware-export/courseware2-work/reused-pdf-crops/"
    "slide09-packshot.png"
)
ILLUSTRATION_PATH = (
    REPO_DIR
    / "assets/component-library/advice-icons/light-diet-dashenlin/"
    "candidates/light-diet-dashenlin-v1.png"
)


def configure_document() -> Document:
    document = Document()
    properties = document.core_properties
    properties.title = "穿心莲内酯滴丸商品培训课件业务真实内容样本"
    properties.subject = "商品培训课件批量生成业务输入范本"
    properties.author = "大参林医药集团"
    properties.comments = (
        "内容来自用户提供的内部课件复刻样本；正式生产仍须使用公司药师、"
        "数据与合规法务确认后的最终原文和授权素材。"
    )
    properties.created = datetime(2026, 7, 30, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 7, 30, tzinfo=timezone.utc)

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_ASCII
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, before, after in (
        ("Heading 1", 16, 12, 8),
        ("Heading 2", 12.5, 9, 5),
        ("Heading 3", 11, 6, 3),
    ):
        style = styles[style_name]
        style.font.name = FONT_ASCII
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(DEEP_GREEN)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    ignore_style = styles.add_style("Courseware Ignore", WD_STYLE_TYPE.PARAGRAPH)
    ignore_style.base_style = normal
    hint_style = styles.add_style("填写提示", WD_STYLE_TYPE.PARAGRAPH)
    hint_style.base_style = normal

    header = section.header.paragraphs[0]
    left = header.add_run("大参林｜商品培训课件批量生成")
    set_run_font(left, size=8.5, color=DEEP_GREEN, bold=True)
    header.add_run("\t")
    right = header.add_run("真实内容填写范本")
    set_run_font(right, size=8.5, color=MID_GRAY)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.6))

    add_page_number(section.footer.paragraphs[0])
    return document


def add_bullets(document: Document, items: list[str], *, ignored: bool = False) -> None:
    paragraphs = []
    for item in items:
        paragraph = document.add_paragraph(
            style="Courseware Ignore" if ignored else "List Bullet"
        )
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.16)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item if ignored else "")
        if ignored:
            run.text = f"• {item}"
        elif not paragraph.text:
            paragraph.add_run(item)
        for current_run in paragraph.runs:
            set_run_font(current_run, size=9.7)
        paragraphs.append(paragraph)
    return paragraphs


def add_numbered(document: Document, items: list[str]) -> None:
    paragraphs = []
    for item in items:
        paragraph = document.add_paragraph(item, style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(4)
        for run in paragraph.runs:
            set_run_font(run, size=9.8)
        paragraphs.append(paragraph)
    restart_numbering(document, paragraphs)


def add_source_note(document: Document, text: str) -> None:
    add_ignore_paragraph(
        document,
        f"来源说明｜{text}",
        size=8.2,
        color=MID_GRAY,
        after=3,
    )


def add_cover_and_batch_guide(document: Document) -> None:
    add_ignore_paragraph(
        document,
        "商品培训课件模板｜业务 Word 真实样本",
        size=11,
        color=GREEN,
        bold=True,
        after=5,
    )
    title = document.add_paragraph(style="Courseware Ignore")
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("穿心莲内酯滴丸")
    set_run_font(run, size=25, color=DEEP_GREEN, bold=True)
    subtitle = document.add_paragraph(style="Courseware Ignore")
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("风热证商品培训课件｜业务已填内容范本")
    set_run_font(run, size=14, color=MID_GRAY, bold=True)

    add_hint(
        document,
        "这是一份“照着改就能用”的真实内容样本。绿色填写说明不会进入课件；"
        "正式提交时，只替换正文和图片即可。医学、功效、用法、数据和销售话术必须使用公司审核终稿。",
    )

    add_ignore_paragraph(
        document, "3 分钟完成一个商品", size=14, color=DEEP_GREEN, bold=True, after=5
    )
    add_numbered(
        document,
        [
            "复制本 Word，文件名改成“商品名称_版本日期.docx”。",
            "保留栏目标题，把本样本正文替换成新商品的审核内容；没有的栏目直接留空。",
            "真实包装、说明书和证据图直接粘贴授权原图；需要辅助插画时只填写图片主题。",
            "完成后把多份 Word 放进同一个文件夹，整批交给课件项目生成 PPTX/PDF。",
        ],
    )

    add_ignore_paragraph(
        document, "批量文件夹示例", size=12, color=DEEP_GREEN, bold=True, after=4
    )
    folder_table = document.add_table(rows=4, cols=2)
    rows = [
        ("01", "穿心莲内酯滴丸_20260730.docx"),
        ("02", "商品B_20260730.docx"),
        ("03", "商品C_20260730.docx"),
        ("输出", "系统逐份生成同品牌模板 PPTX/PDF；失败文件单独返修"),
    ]
    for index, (label, value) in enumerate(rows):
        folder_table.cell(index, 0).text = label
        folder_table.cell(index, 1).text = value
        set_cell_fill(folder_table.cell(index, 0), LIGHT_GREEN)
    apply_table_geometry(folder_table, [1250, 8110])
    format_table_text(folder_table, body_size=9.3)

    add_ignore_paragraph(
        document,
        "模板规则｜公司封面/封底锁定；缺失栏目留空；内容超出一页时自动复制同版式续页；"
        "业务不需要设置字体、字号、颜色或页面坐标。",
        size=9,
        color=DEEP_GREEN,
        bold=True,
        after=0,
    )


def add_course_and_disease(document: Document) -> None:
    document.add_page_break()
    document.add_heading("课程基本信息", level=1)
    add_hint(
        document,
        "课程标题和章节数量可以调整。下列内容来自本次内部课件样本，"
        "用于展示真实整理方式，不代表跳过正式药学与合规审核。",
    )
    table = document.add_table(rows=5, cols=2)
    rows = [
        ("课程标题", "清热泻火，专攻风热诸证"),
        ("疾病主题", "风热证"),
        ("主推商品", "穿心莲内酯滴丸"),
        ("培训对象", "门店员工（内部学习）"),
        ("一句话导语", "精准辨证是中医临床的基石。"),
    ]
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
        set_cell_fill(table.cell(index, 0), LIGHT_GREEN)
    apply_table_geometry(table, [2200, 7160])
    format_table_text(table, body_size=9.5)

    document.add_heading("一、疾病篇", level=1)
    document.add_heading("什么是风热证？", level=2)
    add_data_paragraph(
        document,
        "定义",
        "指风热之邪侵袭人体肌表，导致卫气被遏、肺失宣肃所表现出的证候。"
        "通俗而言，即身体受到“热风”侵袭，引发发热、咽痛、口干等一系列“热”性症状，"
        "是中医外感表证中常见的证型之一。",
    )
    add_data_paragraph(
        document,
        "主要辨证要点",
        "发热重、恶寒轻、咽痛、口渴、苔薄黄。",
    )
    document.add_heading("治疗原则与警示", level=2)
    add_bullets(
        document,
        [
            "核心原则：辛凉解表、清热解毒。",
            "课件原文警示：风热证为“阳热”之证，绝对禁用麻黄、桂枝等辛温发汗药。",
            "兼证处理：临床多见风热夹湿等兼证，需辨明主次兼证。",
        ],
    )
    add_source_note(document, "内部课件第 1.1、1.4 页；正式课件以公司药师审核终稿为准。")


def add_product_information(document: Document) -> None:
    document.add_page_break()
    document.add_heading("二、商品介绍", level=1)
    add_hint(
        document,
        "这一页直接放药品基础资料。标签可以增删；没有可靠资料的字段留空，系统不会补写。",
    )
    layout = document.add_table(rows=1, cols=2)
    left = layout.cell(0, 0)
    right = layout.cell(0, 1)
    picture_paragraph = left.paragraphs[0]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = picture_paragraph.add_run().add_picture(
        str(PACKSHOT_PATH), width=Inches(2.0)
    )
    picture._inline.docPr.set(
        "descr", "穿心莲内酯滴丸包装，裁自用户提供的内部 PDF"
    )
    caption = left.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run("内部 PDF 复用素材｜正式生产优先替换高清授权原图")
    set_run_font(caption_run, size=7.2, color=MID_GRAY)

    product_rows = [
        ("商品名称", "穿心莲内酯滴丸"),
        ("成份", "穿心莲内酯（纯度高达97%以上）；辅料为聚乙二醇、薄膜包衣预混剂。"),
        ("性状", "黄色包衣滴丸，除去包衣后显类白色；味苦。"),
        ("功能主治", "清热解毒，抗菌消炎。用于上呼吸道感染风热证所致的咽痛。"),
        ("规格", "每袋含穿心莲内酯0.15g。"),
        ("用法用量", "口服，一次1袋，一日3次。（课件原文另含首次服用建议，须经内部复核）"),
        ("禁忌", "对本品过敏者禁用。"),
    ]
    for label, value in product_rows:
        paragraph = right.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=8.6, color=DEEP_GREEN, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=8.6)
    apply_table_geometry(layout, [3000, 6360])
    format_table_text(layout, body_size=8.6)
    add_source_note(document, "商品文字与包装图来自内部课件第 2.1 页；不得据此替代说明书或正式审核。")


def add_selling_points(document: Document) -> None:
    document.add_page_break()
    document.add_heading("三、核心卖点", level=1)
    add_hint(
        document,
        "一条卖点写一行。条数少就少展示，条数多则自动生成同版式续页；"
        "关键数据必须保留依据，并经内部复核。",
    )
    headers = ["卖点分组", "卖点名称", "支撑内容"]
    rows = [
        ("剂型优势", "全国独家", "全国唯一穿心莲内酯滴丸剂型"),
        ("剂型优势", "起效更快", "5–10分钟内全部崩解"),
        ("剂型优势", "易携带", "独立包装，方便携带"),
        ("工艺优势", "纯度高", "穿心莲内酯纯度高达97%"),
        ("工艺优势", "稳定性强", "高速滴丸机，含量均匀度高"),
        ("工艺优势", "安全保证", "0糖0添加，纯中药制剂"),
        ("组方优势", "植物单体制剂", "经典老药，植物单体制剂"),
        ("组方优势", "功能全面", "清热解毒＋抗菌消炎"),
        ("组方优势", "疗效显著", "课件原文含总有效率95%表述，正式使用前须补齐依据并审核"),
    ]
    table = document.add_table(rows=1 + len(rows), cols=3)
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    for row_index, values in enumerate(rows, start=1):
        for col, value in enumerate(values):
            table.cell(row_index, col).text = value
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [1750, 1900, 5710])
    format_table_text(table, header=True, body_size=8.5)
    add_source_note(document, "内部课件第 2.3 页；涉及唯一性、纯度、崩解时间和有效率的数据必须完成证据审核。")


def add_audiences_and_consultation(document: Document) -> None:
    document.add_page_break()
    document.add_heading("四、适宜人群与推荐场景", level=1)
    add_hint(
        document,
        "按“人群—典型场景—需求说明”填写，不需要控制每组字数完全相同。",
    )
    headers = ["人群", "典型场景", "需求说明"]
    rows = [
        ("儿童", "学校感冒交叉感染", "课件原文含送服建议，正式使用前须药师复核"),
        ("中青年", "办公室白领", "办公室空调干燥易咽痛"),
        ("中青年", "商旅人士", "出差出行劳累，抵抗力低"),
        ("中青年", "老师", "课堂授课，易咽炎"),
        ("中青年", "全职妈妈", "换季更新药箱"),
        ("老年", "冬日晨起锻炼", "秋冬呼吸疾病高发，易感冒"),
        ("老年", "北方暖气环境", "空气干燥，提前备药"),
    ]
    table = document.add_table(rows=1 + len(rows), cols=3)
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    for row_index, values in enumerate(rows, start=1):
        for col, value in enumerate(values):
            table.cell(row_index, col).text = value
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [1450, 3050, 4860])
    format_table_text(table, header=True, body_size=8.8)

    document.add_heading("望闻问切沟通框架", level=2)
    add_bullets(
        document,
        [
            "望：观察顾客神态、面色与精神状态，初步判断体质倾向。",
            "闻：耐心倾听身体不适与核心诉求，捕捉深层痛点。",
            "问：询问饮食、作息、既往病史等细节，厘清证型。",
            "切：结合辨证结果，匹配产品与服务并解释益处。",
        ],
    )
    add_source_note(document, "内部课件第 3.1、3.2 页。")


def add_scenarios(document: Document) -> None:
    document.add_page_break()
    document.add_heading("五、场景方案", level=1)
    add_hint(
        document,
        "每个场景按同样的四块填写。新增场景时复制整块；联合用药和话术必须使用审核原文。",
    )
    document.add_heading("场景 1｜顾客主诉“喉咙痛”", level=2)
    scenario_rows = [
        ("辨证沟通", "询问是否觉得身上热，用于区分课件中的风热型与风寒型表述。"),
        ("核心用药", "主推穿心莲内酯滴丸；强调课件中的清热解毒、抗菌消炎及滴丸剂型卖点。"),
        ("关联服务", "课件原文关联熊胆薄荷含片，并给出多饮温水、忌辛辣刺激等建议。"),
        ("服务要点", "辨证精准是前提，剂型优势是卖点，关联用药是增值，健康叮嘱是保障。"),
    ]
    table = document.add_table(rows=len(scenario_rows), cols=2)
    for row_index, (label, value) in enumerate(scenario_rows):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
        set_cell_fill(table.cell(row_index, 0), LIGHT_GREEN)
    apply_table_geometry(table, [1900, 7460])
    format_table_text(table, body_size=9.2)

    document.add_heading("场景 2｜顾客主诉“感冒了，发烧，有黄痰”", level=2)
    scenario_rows = [
        ("症状辨析", "课件样本描述：发烧38℃以上、自觉身热不恶寒、咳嗽伴黄痰。"),
        ("核心推荐", "课件样本推荐穿心莲内酯滴丸，并强调清热解毒、药力集中与起效快。"),
        ("联合用药", "课件原文涉及复方氨酚烷胺片联用；必须经公司药师审核后使用。"),
        ("关键话术", "课件原文含“治标＋治本”销售话术；正式课件应粘贴最终审核版本。"),
    ]
    table = document.add_table(rows=len(scenario_rows), cols=2)
    for row_index, (label, value) in enumerate(scenario_rows):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
        set_cell_fill(table.cell(row_index, 0), LIGHT_GREEN)
    apply_table_geometry(table, [1900, 7460])
    format_table_text(table, body_size=9.1)
    add_source_note(document, "内部课件第 3.3、3.4 页；温度、证型、联合用药和话术均需正式审核。")


def add_daily_care(document: Document) -> None:
    document.add_page_break()
    document.add_heading("六、注意事项与日常关怀", level=1)
    add_hint(
        document,
        "每条建议独立填写。药品禁忌优先来自说明书；生活建议也应经内部审核。",
    )
    add_numbered(
        document,
        [
            "饮食之宜：日常多喝水；课件样本列有菊花茶、绿豆汤及梨、西瓜、苦瓜等。",
            "饮食之忌：避免辛辣刺激、油腻煎炸食物，并慎食课件中列举的温补食材。",
            "保证充足休息：建议保证7–8小时高质量睡眠，避免熬夜。",
            "保持室内通风：课件样本建议每日开窗通风2–3次，每次30分钟。",
            "减少过度用嗓：少说话，避免大声喊叫或长时间交谈，让咽喉和声带休息。",
        ],
    )
    add_source_note(document, "内部课件第 4.1 页；时长、频次和饮食建议需由审核人员确认。")

    document.add_heading("七、其他内容（可留空）", level=1)
    add_hint(
        document,
        "不知道应放在哪个栏目时，把原文放在这里。系统会列入待确认，不会擅自丢弃或改写。",
    )
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("本样本无其他待归类内容。")
    set_run_font(run, size=9.8, color=MID_GRAY, italic=True)


def add_image_guide(document: Document) -> None:
    document.add_page_break()
    add_ignore_paragraph(
        document, "图片准备｜两种方式都支持", size=18, color=DEEP_GREEN, bold=True, after=6
    )
    add_hint(
        document,
        "A. 有授权图片：直接粘贴；B. 没有插画：填写一句“图片主题”，项目自动生成候选或输出提示词。"
        "真实包装、说明书、检测报告、处方和品牌证据禁止 AI 仿造。",
    )

    headers = ["课件位置", "处理方式", "业务提交内容", "直观示例"]
    rows = [
        ("02 商品展示", "授权原图", "粘贴高清商品包装图", ""),
        (
            "04 日常关怀",
            "AI 辅助插画",
            "图片主题：药师提醒成年患者清淡饮食，餐桌上有蔬菜和清淡食物",
            "",
        ),
    ]
    table = document.add_table(rows=1 + len(rows), cols=4)
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    for row_index, values in enumerate(rows, start=1):
        for col, value in enumerate(values):
            table.cell(row_index, col).text = value

    p = table.cell(1, 3).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(PACKSHOT_PATH), width=Inches(0.9))
    shape._inline.docPr.set("descr", "内部 PDF 商品包装裁图")
    p = table.cell(2, 3).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(ILLUSTRATION_PATH), width=Inches(0.9))
    shape._inline.docPr.set("descr", "清淡饮食候选插画，非医学证据图")
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [1500, 1550, 4100, 2210])
    format_table_text(table, header=True, body_size=8.2)

    add_ignore_paragraph(
        document,
        "插画质量要求",
        size=12,
        color=DEEP_GREEN,
        bold=True,
        after=4,
    )
    add_bullets(
        document,
        [
            "浅色背景、品牌绿色主线、柔和扁平医药科普插画；人物与医疗物件自然。",
            "图片本身不生成文字、药名、剂量、Logo、水印或新的医学结论。",
            "建议方形高清 PNG，不低于 1536×1536；四周留安全区，便于课件裁切。",
            "AI 图只作辅助候选，必须经过药师和视觉复核后再进入正式 PPTX。",
        ],
        ignored=True,
    )
    add_source_note(
        document,
        "商品包装示例裁自用户提供的内部 PDF；清淡饮食图为候选风格示意、非医学证据图。",
    )


def add_submission_checklist(document: Document) -> None:
    document.add_page_break()
    add_ignore_paragraph(
        document, "提交前 1 分钟自查", size=18, color=DEEP_GREEN, bold=True, after=7
    )
    items = [
        "一份 Word 只写一个商品；批量时复制成多份文件。",
        "商品名称和课程主题已替换；没有资料的栏目已留空，而不是自行补写。",
        "功效、用法用量、禁忌、数据、联合用药和销售话术均为内部审核终稿。",
        "真实包装、说明书和证据资料均使用公司授权原图。",
        "需要 AI 插画的位置只填写了图片主题，且没有要求生成药名、剂量、Logo 或包装。",
        "文件名使用“商品名称_版本日期.docx”，多份 Word 已放入同一文件夹。",
    ]
    for item in items:
        add_ignore_paragraph(document, f"□ {item}", size=10.2, color=TEXT, after=5)

    add_hint(
        document,
        "系统会自动识别栏目、生成同品牌版式，并把缺失项保留为空；"
        "生成后的 PPTX 仍可继续补文字和粘贴最终图片。正式发布前必须完成内容与素材审核。",
    )

    add_ignore_paragraph(
        document,
        "这份真实样本怎么复用",
        size=13,
        color=DEEP_GREEN,
        bold=True,
        after=4,
    )
    add_numbered(
        document,
        [
            "保留各级栏目标题。",
            "删除本商品正文，粘贴新商品的审核内容。",
            "按需要增删条目；不要为了凑页数编写内容。",
            "保存为新文件，继续复制下一商品。",
        ],
    )
    add_ignore_paragraph(
        document,
        "合规声明｜本样本内容来自用户提供的内部课件，仅用于展示业务资料整理和模板生成方法；"
        "它不替代药品说明书、公司药师审核、数据证据审核或合规法务确认。",
        size=8.5,
        color=MID_GRAY,
        after=0,
    )


def build_document() -> Document:
    document = configure_document()
    add_cover_and_batch_guide(document)
    add_course_and_disease(document)
    add_product_information(document)
    add_selling_points(document)
    add_audiences_and_consultation(document)
    add_scenarios(document)
    add_daily_care(document)
    add_image_guide(document)
    add_submission_checklist(document)
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
