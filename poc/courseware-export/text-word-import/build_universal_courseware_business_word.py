#!/usr/bin/env python3
"""Build the universal notebook-like business Word for adaptive courseware."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_DIR = Path("/Users/liminrong/Projects/chain-pharmacy-content-studio")
OUTPUT = (
    REPO_DIR
    / "outputs/courseware-natural-import/培训课件内容与素材提交_通用模板.docx"
)

FONT = "PingFang SC"
BLACK = RGBColor(26, 32, 38)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(96, 107, 118)


def set_run_font(run, *, size: float, color=BLACK, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_free_page_number(paragraph):
    paragraph.alignment = 2
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])
    set_run_font(run, size=9, color=MUTED)


def add_label_line(document, label: str, placeholder: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(f"{label}："), size=11, bold=True)
    set_run_font(paragraph.add_run(placeholder), size=11, color=MUTED)


def add_section_block(document):
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("板块标题（请替换）")

    body = document.add_paragraph()
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.25
    set_run_font(
        body.add_run(
            "在这里直接写本板块的审核内容。可以分段，也可以列出若干要点；"
            "有几条就写几条，不需要凑数量。"
        ),
        size=11,
        color=MUTED,
        italic=True,
    )

    for _ in range(2):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        set_run_font(spacer.add_run(" "), size=11)

    image = document.add_paragraph()
    image.paragraph_format.space_before = Pt(4)
    image.paragraph_format.space_after = Pt(4)
    set_run_font(
        image.add_run("【在此处直接粘贴与本板块相关的图片；没有图片可删除本行】"),
        size=10,
        color=MUTED,
        italic=True,
    )

    caption = document.add_paragraph()
    caption.paragraph_format.space_after = Pt(14)
    set_run_font(caption.add_run("图片说明／来源（可不填）："), size=10, bold=True)
    set_run_font(caption.add_run(" "), size=10, color=MUTED)


def create_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading1 = document.styles["Heading 1"]
    heading1.font.name = FONT
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    heading1.font.size = Pt(16)
    heading1.font.color.rgb = BLUE
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.keep_with_next = True

    ignore = document.styles.add_style("Universal Ignore", WD_STYLE_TYPE.PARAGRAPH)
    ignore.base_style = normal
    ignore.font.name = FONT
    ignore._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    ignore.font.size = Pt(9.5)
    ignore.font.color.rgb = MUTED
    ignore.paragraph_format.space_after = Pt(5)
    ignore.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    set_run_font(
        footer.add_run("一份 Word 对应一个课件主题｜"),
        size=9,
        color=MUTED,
    )
    set_cell_free_page_number(footer)
    return document


def build() -> Document:
    document = create_document()
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(
        title.add_run("培训课件内容与素材提交"),
        size=24,
        color=BLACK,
        bold=True,
    )

    subtitle = document.add_paragraph(style="Universal Ignore")
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(
        subtitle.add_run("像写记事本一样整理内容，AI 负责选页型、排版和拆页。"),
        size=11,
        color=BLUE,
        bold=True,
    )

    guide = document.add_paragraph(style="Universal Ignore")
    set_run_font(
        guide.add_run(
            "填写方法：写主题 → 按自然逻辑写板块 → 图片直接粘贴在相关板块下面。"
            "板块可自由增加或删除，不需要填写页码、模板、卡片数量或动画要求。"
        ),
        size=9.5,
        color=MUTED,
    )

    document.add_paragraph()
    add_label_line(document, "课件主题", "请填写")
    add_label_line(document, "培训对象（可不填）", "请填写")
    add_label_line(document, "培训目标（可不填）", "请填写")

    add_section_block(document)

    more = document.add_paragraph(style="Universal Ignore")
    more.paragraph_format.space_before = Pt(8)
    set_run_font(
        more.add_run(
            "需要更多板块时，复制上面的“板块标题＋内容＋图片”整段继续填写；"
            "不需要的板块直接删除。"
        ),
        size=9.5,
        color=MUTED,
        italic=True,
    )

    document.core_properties.title = "培训课件内容与素材提交通用模板"
    document.core_properties.subject = "多主题培训课件自适应生成业务输入"
    document.core_properties.keywords = (
        "培训课件, Word, 内容提交, 图片粘贴, 自适应页型, 批量创建"
    )
    return document


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
