#!/usr/bin/env python3
"""Build the universal video-training Word and two real filled examples."""

from __future__ import annotations

import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
COURSEWARE_WORD_DIR = ROOT / "poc/courseware-export/text-word-import"
sys.path.insert(0, str(COURSEWARE_WORD_DIR))

from build_universal_courseware_business_word import (  # noqa: E402
    BLACK,
    BLUE,
    MUTED,
    create_document,
    set_run_font,
)


OUTPUT_DIR = ROOT / "outputs/video-training-natural-import"
BLANK_OUTPUT = OUTPUT_DIR / "视频培训内容与素材提交_通用模板.docx"
HEALTH_OUTPUT = OUTPUT_DIR / "风热证健康知识视频培训_真实已填样本.docx"
PRODUCT_OUTPUT = OUTPUT_DIR / "辅酶Q10商品培训视频_真实已填样本.docx"

FEVER_IMAGE = ROOT / "assets/component-library/symptoms/fever/master/fever-v1.png"
THROAT_IMAGE = (
    ROOT / "assets/component-library/symptoms/sore-throat/master/sore-throat-v1.png"
)
GENERIC_PRODUCT_IMAGE = (
    ROOT
    / "assets/component-library/products/generic-coq10/transparent/"
    "generic-coq10-packshot-v1.png"
)


def add_title_block(document, subtitle: str, note: str):
    title = document.add_paragraph()
    title.paragraph_format.space_after = Inches(0.04)
    set_run_font(
        title.add_run("视频培训内容与素材提交"),
        size=24,
        color=BLACK,
        bold=True,
    )

    subtitle_paragraph = document.add_paragraph(style="Universal Ignore")
    set_run_font(
        subtitle_paragraph.add_run(subtitle),
        size=11,
        color=BLUE,
        bold=True,
    )

    note_paragraph = document.add_paragraph(style="Universal Ignore")
    set_run_font(
        note_paragraph.add_run(note),
        size=9.5,
        color=MUTED,
    )


def add_meta(document, label: str, value: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Inches(0.07)
    set_run_font(paragraph.add_run(f"{label}："), size=11, bold=True)
    set_run_font(paragraph.add_run(value), size=11)


def add_section(document, title: str, paragraphs: list[str]):
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(title)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        set_run_font(paragraph.add_run(text), size=11)


def add_images(
    document,
    images: list[tuple[Path, str]],
    caption: str,
    *,
    width: float,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    for path, alt in images:
        run = paragraph.add_run()
        picture = run.add_picture(str(path), width=Inches(width))
        picture._inline.docPr.set("descr", alt)
        picture._inline.docPr.set("title", alt)
        paragraph.add_run("  ")

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        caption_paragraph.add_run(f"图片说明／来源：{caption}"),
        size=9,
        color=MUTED,
        italic=True,
    )


def add_properties(document, title: str, subject: str, keywords: str):
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.keywords = keywords


def build_blank():
    document = create_document()
    add_title_block(
        document,
        "像写记事本一样整理内容，AI 负责拆章节、选画面和安排时长。",
        (
            "填写方法：写主题 → 按自然逻辑写板块 → 图片直接粘贴在相关板块下面。"
            "板块可自由增加、删除或调整顺序，不需要照着现有视频章节填写。"
        ),
    )
    document.add_paragraph()
    add_meta(document, "课件主题", "请填写")
    add_meta(document, "培训对象（可不填）", "请填写")
    add_meta(document, "培训目标（可不填）", "请填写")
    add_meta(document, "期望时长（可不填）", "请填写")

    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("板块标题（请替换）")

    body = document.add_paragraph(style="Universal Ignore")
    set_run_font(
        body.add_run(
            "在这里直接写本板块的审核内容。正文应当可以直接用于讲解或旁白；"
            "可以分段，也可以列出若干要点，有几条就写几条。"
        ),
        size=11,
        color=MUTED,
        italic=True,
    )
    for _ in range(2):
        spacer = document.add_paragraph()
        set_run_font(spacer.add_run(" "), size=11)

    image = document.add_paragraph(style="Universal Ignore")
    set_run_font(
        image.add_run("【在此处直接粘贴与本板块相关的图片；没有图片可删除本行】"),
        size=10,
        color=MUTED,
        italic=True,
    )
    caption = document.add_paragraph()
    set_run_font(caption.add_run("图片说明／来源（可不填）："), size=10, bold=True)

    more = document.add_paragraph(style="Universal Ignore")
    more.paragraph_format.space_before = Inches(0.12)
    set_run_font(
        more.add_run(
            "需要更多板块时，复制上面的“板块标题＋正文＋图片”整段；"
            "一份 Word 对应一个视频主题。"
        ),
        size=9.5,
        color=MUTED,
        italic=True,
    )
    add_properties(
        document,
        "视频培训内容与素材提交通用模板",
        "健康知识与商品培训视频的内容驱动业务输入",
        "视频培训, Word, 内容提交, 图片粘贴, 自适应章节, 批量创建",
    )
    return document


def build_health_sample():
    document = create_document()
    add_title_block(
        document,
        "真实已填样本｜风热证健康知识培训",
        (
            "本样本只示范如何整理自然板块、讲解正文和相关图片；"
            "新项目不要求使用相同章节，投产前仍需完成医学、药事、合规和素材授权审核。"
        ),
    )
    add_meta(document, "课件主题", "中医基础知识——风热证")
    add_meta(document, "培训对象（可不填）", "连锁药店一线员工")
    add_meta(
        document,
        "培训目标（可不填）",
        "理解风热证的基本概念、典型表现、调理思路和日常注意事项。",
    )
    add_meta(document, "期望时长（可不填）", "约 2—3 分钟")

    add_section(
        document,
        "问题引入",
        [
            "喉咙又肿又痛，咳嗽时有黄痰，鼻涕又黄又稠，身体发热又口渴——这些偏“热”的表现，可能会让人想到风热证。"
        ],
    )
    add_section(
        document,
        "基本概念与典型表现",
        [
            "简单来说，风热证就是风邪和热邪一起侵入身体，导致体表不适、肺气不顺。",
            "常见表现可以包括发热、口渴、咽喉肿痛、咳嗽、痰黄或黄稠鼻涕等。具体判断仍应结合完整情况，由专业人员评估。",
        ],
    )
    add_images(
        document,
        [
            (FEVER_IMAGE, "发热症状插画"),
            (THROAT_IMAGE, "喉咙肿痛症状插画"),
        ],
        "从左至右：发热、喉咙肿痛；项目内原创症状插画",
        width=2.0,
    )
    add_section(
        document,
        "调理思路",
        [
            "课程中的核心调理思路是疏风清热。",
            "涉及草本、药品、剂量、用法或搭配建议时，应完整使用公司审核终稿，不根据样本自行补充。",
        ],
    )
    add_section(
        document,
        "日常注意事项",
        [
            "房间适当开窗通风，保持空气流通。",
            "少量多次饮用温水，饮食以清淡为主，并保证充足休息。",
            "出现持续高热、症状加重或其他异常情况时，应及时寻求专业帮助。",
        ],
    )
    add_section(
        document,
        "课程小结",
        [
            "这次培训重点理解三个方面：风热证的基本概念、偏热的典型表现，以及疏风清热和日常护理的基本思路。"
        ],
    )
    add_properties(
        document,
        "风热证健康知识视频培训真实已填样本",
        "通用视频内容 Word 的健康类真实填写参考",
        "健康知识视频, 风热证, 真实样本, 内容驱动, 授权图片",
    )
    return document


def build_product_sample():
    document = create_document()
    add_title_block(
        document,
        "真实已填样本｜辅酶 Q10 商品培训",
        (
            "本样本只示范如何整理自然板块和讲解正文；章节不是固定模板。"
            "包装图为无品牌示意，正式生产必须替换为公司授权包装及审核内容。"
        ),
    )
    add_meta(document, "课件主题", "辅酶 Q10 商品培训")
    add_meta(document, "培训对象（可不填）", "连锁药店门店员工")
    add_meta(
        document,
        "培训目标（可不填）",
        "理解审核稿中的商品定位、基础信息、核心知识和适用沟通边界。",
    )
    add_meta(document, "期望时长（可不填）", "系统根据审核正文估算")

    add_section(
        document,
        "为什么要了解辅酶 Q10",
        [
            "心脏持续工作需要稳定的能量供应。了解辅酶 Q10 与细胞能量生成的关系，有助于进一步理解这类商品知识。"
        ],
    )
    add_section(
        document,
        "商品基础信息",
        [
            "本次培训商品为辅酶 Q10 胶囊。正式课程中的商品名称、剂型、规格、批准文号和包装文字，以公司审核资料和授权包装原图为准。"
        ],
    )
    add_images(
        document,
        [(GENERIC_PRODUCT_IMAGE, "无品牌辅酶 Q10 包装示意")],
        "项目内无品牌包装示意，仅展示图片粘贴位置；正式生产必须替换为公司授权包装原图",
        width=2.6,
    )
    add_section(
        document,
        "核心知识",
        [
            "辅酶 Q10 存在于细胞线粒体内，参与细胞能量生成过程。涉及功效、抗氧化或其他专业表述时，应逐句采用公司最终审核稿。",
            "工艺、原料、检测数据和证据内容只有在公司确认可用于内部培训后，才进入视频。",
        ],
    )
    add_section(
        document,
        "适宜人群与联合方案",
        [
            "适宜人群、咨询边界和沟通话术以公司审核稿为准，不能仅根据日常不适自行判断。",
            "联合方案应完整提交药品名称、搭配逻辑和获准话术；实际用药应遵医嘱或咨询专业人员。",
        ],
    )
    add_section(
        document,
        "课程小结",
        [
            "本课程从辅酶 Q10 与细胞能量生成的关系出发，介绍商品基础信息、核心知识、适用沟通边界和联合方案要求。"
        ],
    )
    add_properties(
        document,
        "辅酶Q10商品培训视频真实已填样本",
        "通用视频内容 Word 的商品类真实填写参考",
        "商品培训视频, 辅酶Q10, 真实样本, 内容驱动, 包装授权",
    )
    return document


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [
        (build_blank(), BLANK_OUTPUT),
        (build_health_sample(), HEALTH_OUTPUT),
        (build_product_sample(), PRODUCT_OUTPUT),
    ]
    for document, path in outputs:
        document.save(path)
        print(path)


if __name__ == "__main__":
    main()
