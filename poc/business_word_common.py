"""Shared visual master for business-facing training-video Word samples."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# compact_reference_guide + customer_pack named overrides
FONT = "Source Han Sans SC"
NAVY = "123B54"
BLUE = "217AB7"
TEAL = "168695"
TEXT = "243642"
MUTED = "617582"
WHITE = "FFFFFF"
LIGHT_BLUE = "EAF5FC"
LIGHT_TEAL = "E9F6F7"
LIGHT_GOLD = "F8F0D9"
BORDER = "C9DCE6"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run(run, *, size=11, color=TEXT, bold=False, italic=False):
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        fonts.set(qn(key), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run(run, **kwargs)
    return run


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        normal._element.rPr.rFonts.set(qn(key), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Title", 27, 0, 8, NAVY),
        ("Heading 1", 16, 18, 10, NAVY),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, NAVY),
    ):
        style = styles[name]
        style.font.name = FONT
        for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            style._element.rPr.rFonts.set(qn(key), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = styles["List Bullet"]
    list_style.base_style = normal
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = Inches(-0.188)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    if "Courseware Ignore" not in styles:
        ignore = styles.add_style("Courseware Ignore", WD_STYLE_TYPE.PARAGRAPH)
        ignore.base_style = normal


def create_document(document_type):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    _configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(
        header,
        f"视频培训业务内容输入｜{document_type}",
        size=8.5,
        color=MUTED,
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(
        footer,
        "内部培训资料 · 仅使用已审核内容和授权素材　",
        size=8,
        color=MUTED,
    )
    page = OxmlElement("w:fldSimple")
    page.set(qn("w:instr"), "PAGE")
    footer._p.append(page)
    return document


def add_body(
    document,
    text,
    *,
    size=11,
    color=TEXT,
    bold=False,
    italic=False,
    before=0,
    after=6,
    align=None,
    style=None,
):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    add_text(
        paragraph,
        text,
        size=size,
        color=color,
        bold=bold,
        italic=italic,
    )
    return paragraph


def add_callout(document, text, *, fill=LIGHT_TEAL, color=NAVY, bold=False):
    paragraph = document.add_paragraph(style="Courseware Ignore")
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.25
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)
    add_text(paragraph, text, size=10.5, color=color, bold=bold)
    return paragraph


def add_bullets(document, items, *, style="Courseware Ignore", color=TEXT):
    if not items:
        items = ["无"]
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        add_text(paragraph, item, size=10.5, color=color)
        if style == "Courseware Ignore":
            paragraph.style = document.styles["Courseware Ignore"]
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph._p.get_or_add_pPr().insert(0, _bullet_num_pr(document))


def _bullet_num_pr(document):
    source = document.styles["List Bullet"].element.pPr.numPr
    if source is None:
        return OxmlElement("w:numPr")
    return deepcopy(source)


def add_key_value_table(document, pairs):
    table = document.add_table(rows=len(pairs) + 1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for index, text in enumerate(("课程信息", "业务填写内容")):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_text(paragraph, text, size=9, color=WHITE, bold=True)
    for row, (label, value) in zip(table.rows[1:], pairs):
        _set_cell_shading(row.cells[0], LIGHT_BLUE)
        label_p = row.cells[0].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        add_text(label_p, label, size=9.5, color=NAVY, bold=True)
        value_p = row.cells[1].paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        add_text(value_p, value or "无", size=9.5)
    return table


def add_business_fields(document, screen_text, assets):
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for row, label, items in zip(
        table.rows,
        ("必须原样上屏的事实／短文案", "本章节授权素材"),
        (screen_text, assets),
    ):
        label_cell, value_cell = row.cells
        _set_cell_shading(label_cell, LIGHT_BLUE)
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        add_text(label_p, label, size=9.5, color=NAVY, bold=True)
        for index, item in enumerate(items or ["无"]):
            paragraph = value_cell.paragraphs[0] if index == 0 else value_cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.2
            add_text(paragraph, f"• {item}", size=9.5)
    return table


def add_cover(
    document,
    *,
    document_type,
    example_title,
    notice,
    showcase_frames,
):
    add_body(
        document,
        "视频培训业务内容输入",
        size=11,
        color=TEAL,
        bold=True,
        after=2,
        style="Courseware Ignore",
    )
    add_body(
        document,
        f"{document_type}\n业务 Word 填写示例",
        size=27,
        color=NAVY,
        bold=True,
        after=7,
        style="Courseware Ignore",
    )
    add_body(
        document,
        f"真实已填示例：{example_title}",
        size=11,
        color=BLUE,
        bold=True,
        after=10,
        style="Courseware Ignore",
    )
    add_callout(document, notice, fill=LIGHT_GOLD, bold=True)
    add_body(
        document,
        "业务只填写：①审核原文　②必须原样上屏的事实／短文案　③授权素材",
        size=10.5,
        color=NAVY,
        bold=True,
        after=5,
        style="Courseware Ignore",
    )
    add_body(
        document,
        "章节可以增删、复制、重排；换主题时替换内容和授权素材。画面配方、动效、镜头时码、素材匹配与缺口由系统生成。",
        size=10,
        color=MUTED,
        after=9,
        style="Courseware Ignore",
    )
    document.add_page_break()
    add_body(
        document,
        "沉淀目标模板真实画面",
        size=12,
        color=NAVY,
        bold=True,
        after=4,
        style="Courseware Ignore",
    )
    add_body(
        document,
        "以下画面直接截自本项目指定的沉淀目标模板，不是另外生成的示意图。"
        "新主题会替换标题、正文、图片和商品素材；母版、字体、角色、字幕与动效语言保持一致。",
        size=9,
        color=MUTED,
        after=5,
        style="Courseware Ignore",
    )
    if len(showcase_frames) != 3:
        raise ValueError("showcase_frames must contain exactly three real video frames")
    lead = showcase_frames[0]
    picture = document.add_picture(str(lead["path"]), width=Inches(6.45))
    picture_paragraph = document.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_after = Pt(2)
    doc_pr = picture._inline.docPr
    doc_pr.set("descr", f"{document_type}沉淀模板原帧：{lead['label']}，{lead['source_mark']}")
    doc_pr.set("title", f"沉淀模板原帧：{lead['label']}")
    add_body(
        document,
        f"{lead['label']}｜{lead['source_mark']}",
        size=8.5,
        color=MUTED,
        after=5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        style="Courseware Ignore",
    )
    comparison = document.add_table(rows=1, cols=2)
    set_table_geometry(comparison, [4680, 4680])
    for cell, frame in zip(comparison.rows[0].cells, showcase_frames[1:]):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run()
        image = run.add_picture(str(frame["path"]), width=Inches(2.92))
        image._inline.docPr.set(
            "descr",
            f"{document_type}沉淀模板原帧：{frame['label']}，{frame['source_mark']}",
        )
        image._inline.docPr.set("title", f"沉淀模板原帧：{frame['label']}")
        caption = cell.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(0)
        add_text(
            caption,
            f"{frame['label']}｜{frame['source_mark']}",
            size=8.2,
            color=MUTED,
        )
    add_body(
        document,
        "模板截图只用于说明真实课件外观；新项目不直接复用示例中的文字、图片、Logo 或商品像素。",
        size=8.5,
        color=MUTED,
        italic=True,
        before=4,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        style="Courseware Ignore",
    )


def add_course_info_page(document, pairs):
    document.add_page_break()
    document.add_heading("课程基本信息", level=1)
    add_key_value_table(document, pairs)
    add_callout(
        document,
        "下方唯一的“标题 1”是课程标题；每个“标题 2”是一章。新增、删除或调整章节时保持这两级标题结构。",
    )


def add_chapter(
    document,
    *,
    heading,
    approved_text,
    screen_text,
    assets,
    page_break=True,
):
    if page_break:
        document.add_page_break()
    document.add_heading(heading, level=2)
    add_body(
        document,
        "审核原文",
        size=9.5,
        color=TEAL,
        bold=True,
        after=3,
        style="Courseware Ignore",
    )
    add_body(document, approved_text, size=10.5, after=9)
    add_business_fields(document, screen_text, assets)


def add_reusable_block(document, *, example_hint):
    document.add_page_break()
    document.add_heading("可复制的章节填写块", level=1)
    add_callout(
        document,
        "复制下方整块到课程正文中；章节名称保持“标题 2”。需要几章就复制几份，系统会按新内容重新匹配画面和素材。",
        bold=True,
    )
    add_body(
        document,
        f"章节标题（Heading 2）：{example_hint}",
        size=13,
        color=BLUE,
        bold=True,
        after=7,
        style="Courseware Ignore",
    )
    for label, hint in (
        ("审核原文", "粘贴可直接配音的最终审核原文，不写画面设计和制作要求。"),
        (
            "必须原样上屏的事实／短文案",
            "填写不可改写的名称、规格、数据、剂量、警示语或短句。",
        ),
        (
            "本章节授权素材",
            "填写随 Word 一起提交的文件名、来源和可用范围；没有则写“无”。",
        ),
    ):
        add_body(
            document,
            label,
            size=9.5,
            color=NAVY,
            bold=True,
            after=2,
            style="Courseware Ignore",
        )
        add_callout(document, hint, fill=LIGHT_BLUE, color=MUTED)
    add_callout(
        document,
        "不要填写：画面意图、组件名称、动效、镜头时码、坐标、字体、颜色、图片提示词、逐章节审核状态或禁用画面。",
        fill=LIGHT_GOLD,
        bold=True,
    )


def save_document(document, output: Path, *, title, subject, keywords):
    output.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.author = "Chain Pharmacy Content Studio"
    document.core_properties.keywords = keywords
    document.save(output)
