from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "可可康绿色视频金样_逐镜内容审核确认表_v1.docx"

GREEN = "087C45"
GREEN_DARK = "075C38"
PALE = "EDF6F0"
PALE_2 = "F6FAF7"
GRAY = "66746E"
LIGHT_GRAY = "F2F4F3"
LINE = "D7E2DB"
RED = "9B1C1C"
RED_PALE = "FFF0EC"
BLACK = "17241F"
WHITE = "FFFFFF"


SCENES = [
    {
        "id": "K01", "title": "产品与课程开场", "recipe": "F10 品牌封底／产品主视觉", "source": "refs/04；主题登记",
        "task": "只回答“这是什么产品、这堂课学什么”，建立可可康灵芝胶囊与内部培训身份。",
        "legacy": "商品名为可可康灵芝胶囊；包装可见 OTC、60 粒／瓶。课程可覆盖成分、产品方向、产品特点、联合应用与服用周期。",
        "risk": "授权高清包装尚未提供。审核前不要在开场加入功效承诺。",
        "assets": "可可康高清正面／多面包装；大参林授权 Logo。",
    },
    {
        "id": "K02", "title": "三类状态总览", "recipe": "F01 问题模型 + F03 总览", "source": "refs/03、09",
        "task": "用三个生活状态建立学习问题：经常失眠、常喝酒／肝功能差、免疫力低下。",
        "legacy": "旧课件把三类人群分别写为经常失眠人群、常喝酒伤肝／肝功能差人群、免疫力低下人群。",
        "risk": "请确认最终允许的人群称谓；不得把年龄形象解释为新增适用年龄。",
        "assets": "三类成人生活情境可由制作侧重绘；不需要品牌包装。",
    },
    {
        "id": "K03", "title": "失眠状态", "recipe": "F02 问题形成", "source": "refs/03、05",
        "task": "按时间顺序表现入睡困难、易醒、早醒和次日疲倦，再进入产品知识解释。",
        "legacy": "联合用药页出现“入睡困难、易醒、早醒和醒后再入睡困难”；旧分镜另加入日间疲倦作为生活表现。",
        "risk": "请确认可以上屏的状态词，以及“日间疲倦”是否保留；不新增疾病诊断。",
        "assets": "卧室、时钟、次日工作状态插画，由制作侧重绘。",
    },
    {
        "id": "K04", "title": "饮酒与肝脏负担", "recipe": "F02 问题形成", "source": "refs/03、05",
        "task": "从饮酒生活情境进入肝脏负担，作为保肝方向的前置问题。",
        "legacy": "旧课件写“常喝酒伤肝人群”“肝功能差人群”，联合用药页出现“肝损伤、病毒性肝炎”。",
        "risk": "“病毒性肝炎”“肝损伤”等表述风险较高，请明确保留、改写或删除；体内变化不得由制作侧猜。",
        "assets": "成人饮酒情境可重绘；肝脏变化须按审核后的机制说明制作。",
    },
    {
        "id": "K05", "title": "免疫力低下状态", "recipe": "F02 问题形成", "source": "refs/03、06",
        "task": "表现容易反复不适、抵抗力差的生活状态，建立免疫方向问题。",
        "legacy": "旧课件写“免疫力低下”“易反复生病、抵抗力差”。旧分镜曾使用季节变化、打喷嚏等非诊断性情境。",
        "risk": "请确认“少生病”“反复生病”等措辞是否允许；不把普通不适演成具体疾病。",
        "assets": "成人生活状态插画可重绘。",
    },
    {
        "id": "K06", "title": "三大产品方向", "recipe": "F03 总览 + F08 结论标签", "source": "refs/03、09",
        "task": "把三类状态对应到三个产品知识方向，作为后续课程导航。",
        "legacy": "宁心安神助睡眠；保肝护肝抗衰老；提升免疫少生病。",
        "risk": "三条均属于功效表达，必须逐条确认最终短标题和旁白口径。",
        "assets": "不需要包装；使用中性睡眠、肝脏、免疫图标。",
    },
    {
        "id": "K07", "title": "产品身份与功能主治", "recipe": "F10 产品主视觉", "source": "refs/04 包装截图",
        "task": "回到真包装，讲清 OTC、规格和获准的功能主治。",
        "legacy": "包装截图可见“宁心安神，健脾和胃。用于失眠健忘，身体虚弱，神经衰弱”；规格显示 60 粒／瓶。",
        "risk": "必须以说明书／批准信息终稿核对，不能把低清截图当权威来源或生产资产。",
        "assets": "说明书／批准信息；授权高清包装。",
    },
    {
        "id": "K08", "title": "两类核心成分", "recipe": "F03 角色总览", "source": "refs/01、09",
        "task": "建立两个贯穿后续章节的知识锚点：灵芝多糖、灵芝三萜。",
        "legacy": "旧课件将灵芝多糖、灵芝三萜称为“核心有效成分”。",
        "risk": "请确认“核心有效成分”称谓是否保留，或改为“课件重点成分／主要成分”。",
        "assets": "灵芝／孢子、两类中性结构符号；不复用 PPT 截图。",
    },
    {
        "id": "K09", "title": "灵芝多糖", "recipe": "F04 路径作用", "source": "refs/01",
        "task": "只讲审核后允许保留的多糖相关作用，不扩写分子通路。",
        "legacy": "旧课件包含“增强免疫系统机能”“降低血压，预防心血管疾病的产生”。",
        "risk": "涉及血压、心血管疾病预防等高风险表述；正式旁白与动画必须等待药师／合规终稿。",
        "assets": "审核后再决定是否需要免疫或血流机制图。",
    },
    {
        "id": "K10", "title": "灵芝三萜", "recipe": "F04 路径作用", "source": "refs/01",
        "task": "只讲审核后允许保留的三萜相关作用，不创作受体、靶点或治疗机制。",
        "legacy": "旧课件包含抗炎、镇痛、镇静、抗衰老、抑制肿瘤细胞，以及促进淋巴细胞增殖、提高相关免疫细胞能力等表述。",
        "risk": "包含抗肿瘤与免疫细胞机制等高风险内容，建议逐句确认或删除。",
        "assets": "审核后再决定允许出现的细胞／过程图。",
    },
    {
        "id": "K11", "title": "成分与三大方向对应", "recipe": "F06 输出变化 + F08 标签", "source": "refs/03",
        "task": "把获准的成分知识分别对应到睡眠、肝脏和免疫方向，完成一次机制小结。",
        "legacy": "旧课件把多糖／三萜与调节神经递质、减轻肝功能损伤、增强肝解毒、促进淋巴细胞增殖等关系相连。",
        "risk": "不得自动把关联表述升级成确定性因果；不得绘制未提供的分子靶点。",
        "assets": "审核终稿确认后，制作侧再重绘三条路径。",
    },
    {
        "id": "K12", "title": "特点一：产地与含量", "recipe": "F08 结论标签", "source": "refs/04",
        "task": "用“产地—原料—数据”三步说明第一项产品特点。",
        "legacy": "安徽大别山产区赤灵芝；每 100g 含多糖 9.13g；灵芝三萜含量高。",
        "risk": "9.13g/100g 与“三萜含量高”需要证据及可展示范围；不能凭 PPT 截图制作检测结论。",
        "assets": "产地授权图／视频；检测报告或公司批准证据。",
    },
    {
        "id": "K13", "title": "特点二：双重提取", "recipe": "F04 路径作用 + F08 标签", "source": "refs/04",
        "task": "按工艺顺序演示第一次提取、第二次浓缩提取与胶囊成形。",
        "legacy": "专利浓缩提取；二次浓缩提取／双重提取；含量高、质地纯；胶囊剂型锁住营养、服用方便、吸收好。",
        "risk": "专利、质地纯、吸收好等表述需证据；不得自行增加吸收率、倍数或实验数据。",
        "assets": "专利／工艺证明；公司可用工艺图或允许制作的示意范围。",
    },
    {
        "id": "K14", "title": "特点三：生产与质量", "recipe": "F08 结论标签", "source": "refs/04",
        "task": "用生产主体、质量检查和封装路径说明第三项产品特点。",
        "legacy": "中山可可康生产；通过国家 GMP 的工厂；采用国内先进生产技术；品质保证。",
        "risk": "GMP、先进技术与品质保证均需公司批准的当前证据和措辞。",
        "assets": "中山可可康工厂／生产线／GMP 授权材料。",
    },
    {
        "id": "K15", "title": "三套联合应用总览", "recipe": "F03 方案总览", "source": "refs/05、06",
        "task": "先只建立三个场景与三套组合的记忆关系，不展开销售话术。",
        "legacy": "失眠：谷维素片 + 灵芝胶囊；肝功能异常：护肝片 + 灵芝胶囊；免疫力低下：转移因子口服溶液 + 灵芝胶囊。",
        "risk": "联合用药属于高风险培训内容，三套组合及适用情境必须由药师／合规逐项确认。",
        "assets": "三款联合商品与可可康授权高清包装。",
    },
    {
        "id": "K16", "title": "三套联合方案解释", "recipe": "F04 路径作用 + F08 标签", "source": "refs/05、06",
        "task": "按“问题—组合—各自作用—注意事项”依次解释三套获准方案。",
        "legacy": "旧课件包含谷维素营养神经、护肝片疏肝理气／降低转氨酶、转移因子调节免疫等组合逻辑，并出现“长期吃比较容易导致依赖”等话术。",
        "risk": "全部属于高风险药学话术；尤其“长期容易依赖”不得未经确认进入正式片。",
        "assets": "授权包装；必要时提供说明书或联合方案依据。",
    },
    {
        "id": "K17", "title": "建议服用周期", "recipe": "F07 时间变化", "source": "refs/09",
        "task": "用月历解释获准的服用周期，不增加剂量、频次或疗程承诺。",
        "legacy": "1 个月为 1 个服用周期，建议连续服用 2–3 个月。",
        "risk": "必须与说明书或审核终稿一致；确认是否需要增加适用条件或注意事项。",
        "assets": "不需要品牌资产；只需审核后的时间数据。",
    },
    {
        "id": "K18", "title": "六维总结与品牌封底", "recipe": "F09 结果回扣 + F10 封底", "source": "refs/09；K01–K17 终稿",
        "task": "从成分、产品方向、人群、特点、联合应用、周期六个维度回到产品，完成内部培训收束。",
        "legacy": "旧总结页列出有效成分、产品卖点、三大核心功效、适应人群和建议服用周期；联合用药由前页补充。",
        "risk": "本镜只能汇总已在前 17 镜通过的内容，不得重新加入被删除或未审核表述。",
        "assets": "授权真包装、大参林 Logo、内部培训声明。",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def font_run(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = "Source Han Sans SC"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Source Han Sans SC")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Source Han Sans SC")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Source Han Sans SC")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=11, bold=False, color=BLACK, italic=False):
    run = paragraph.add_run(text)
    font_run(run, size, bold, color, italic)
    return run


def style_paragraph(paragraph, after=6, before=0, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def set_style_font(style, size, color=BLACK, bold=False, before=0, after=6, line=1.25):
    style.font.name = "Source Han Sans SC"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Source Han Sans SC")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Source Han Sans SC")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Source Han Sans SC")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    add_text(p, text, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=GREEN_DARK)
    return p


def add_label_value(doc, label, value, fill=None, color=BLACK):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1700, 7660])
    set_table_borders(table)
    set_cell_shading(table.cell(0, 0), PALE)
    if fill:
        set_cell_shading(table.cell(0, 1), fill)
    p0 = table.cell(0, 0).paragraphs[0]
    style_paragraph(p0, after=0)
    add_text(p0, label, bold=True, color=GREEN_DARK)
    p1 = table.cell(0, 1).paragraphs[0]
    style_paragraph(p1, after=0)
    add_text(p1, value, color=color)
    return table


def add_scene_block(doc, scene):
    h = doc.add_paragraph()
    style_paragraph(h, before=8, after=5, line=1.1, keep=True)
    add_text(h, f"{scene['id']}  {scene['title']}", size=14, bold=True, color=GREEN_DARK)
    add_text(h, f"    {scene['recipe']}", size=9.5, color=GRAY)

    meta = doc.add_table(rows=1, cols=2)
    set_table_geometry(meta, [4680, 4680])
    set_table_borders(meta)
    set_cell_shading(meta.cell(0, 0), PALE_2)
    set_cell_shading(meta.cell(0, 1), PALE_2)
    p = meta.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, "内容来源：", bold=True, color=GREEN_DARK)
    add_text(p, scene["source"], size=10)
    p = meta.cell(0, 1).paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, "当前结论：", bold=True, color=GREEN_DARK)
    add_text(p, "待业务／药师确认", size=10, color=RED)

    rows = [
        ("教学任务", scene["task"], PALE_2, BLACK),
        ("旧课件口径\n（非审核终稿）", scene["legacy"], WHITE, BLACK),
        ("必须确认", scene["risk"], RED_PALE, RED),
        ("授权／证据依赖", scene["assets"], WHITE, BLACK),
    ]
    for label, value, fill, color in rows:
        add_label_value(doc, label, value, fill, color)

    decision = doc.add_table(rows=1, cols=2)
    set_table_geometry(decision, [1700, 7660])
    set_table_borders(decision)
    set_cell_shading(decision.cell(0, 0), PALE)
    p = decision.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, "审核结论", bold=True, color=GREEN_DARK)
    p = decision.cell(0, 1).paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, "☐ 保留    ☐ 改写    ☐ 删除    ☐ 需补证据／素材", size=10.5)

    final = doc.add_table(rows=1, cols=2)
    set_table_geometry(final, [1700, 7660])
    set_table_borders(final)
    set_cell_shading(final.cell(0, 0), PALE)
    p = final.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, "批准终稿／备注", bold=True, color=GREEN_DARK)
    cell = final.cell(0, 1)
    set_cell_shading(cell, "FBFCFB")
    p = cell.paragraphs[0]
    style_paragraph(p, after=2)
    add_text(p, "【请填写可直接用于旁白与上屏的终稿；删除时写明原因】", size=9.5, color=GRAY, italic=True)
    for _ in range(2):
        p = cell.add_paragraph()
        style_paragraph(p, after=0)
        add_text(p, " ")


def add_footer(section):
    footer = section.footer
    footer.distance = Inches(0.492)
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=0, line=1.0)
    add_text(p, "可可康绿色视频金样 · 内容审核确认表    ", size=8.5, color=GRAY)
    add_text(p, "第 ", size=8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)
    add_text(p, " 页", size=8.5, color=GRAY)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    set_style_font(styles["Normal"], 10.5, BLACK, False, 0, 5, 1.18)
    set_style_font(styles["Heading 1"], 16, GREEN_DARK, True, 18, 10, 1.0)
    set_style_font(styles["Heading 2"], 13, GREEN_DARK, True, 14, 7, 1.0)
    set_style_font(styles["Heading 3"], 12, GREEN_DARK, True, 10, 5, 1.0)

    header = section.header
    hp = header.paragraphs[0]
    style_paragraph(hp, after=0, line=1.0)
    add_text(hp, "大参林内部培训 · 内容审核工作件", size=8.5, bold=True, color=GREEN_DARK)

    # Customer-pack style first page.
    p = doc.add_paragraph()
    style_paragraph(p, before=10, after=2)
    add_text(p, "CONTENT REVIEW PACK", size=9, bold=True, color=GREEN)
    p = doc.add_paragraph()
    style_paragraph(p, after=5, line=1.0)
    add_text(p, "可可康绿色视频金样", size=25, bold=True, color=GREEN_DARK)
    p = doc.add_paragraph()
    style_paragraph(p, after=18, line=1.1)
    add_text(p, "逐镜内容审核确认表 · K01–K18", size=15, color=GRAY)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2400, 6960])
    set_table_borders(table)
    for i, (label, value) in enumerate([
        ("用途", "业务／药师逐镜确认保留、改写、删除与批准终稿"),
        ("内容来源", "公司内部培训 PPT refs/01–06、09；结构参考 Pfizer 案例"),
        ("当前状态", "结构已完成；药学内容、证据与包装授权待确认"),
        ("版本", "v1 · 2026-07-31"),
    ]):
        set_cell_shading(table.cell(i, 0), PALE)
        p0 = table.cell(i, 0).paragraphs[0]
        style_paragraph(p0, after=0)
        add_text(p0, label, bold=True, color=GREEN_DARK)
        p1 = table.cell(i, 1).paragraphs[0]
        style_paragraph(p1, after=0)
        add_text(p1, value)

    p = doc.add_paragraph()
    style_paragraph(p, before=18, after=8)
    add_text(p, "使用规则", size=13, bold=True, color=GREEN_DARK)
    for idx, text in enumerate([
        "逐镜勾选一个主结论；如选择“改写”，必须在批准终稿区写出可直接用于旁白与上屏的文字。",
        "旧课件口径只用于定位来源，不代表已经完成药师／合规审核。",
        "包装、Logo、说明书、专利、GMP 和联合用药证据必须另行提交授权原件。",
        "本表确认内容，不自动确认视觉设计，也不自动授权进入正式生产。",
    ], start=1):
        p = doc.add_paragraph(style="List Number")
        style_paragraph(p, after=5, line=1.18)
        add_text(p, text, size=10.5)

    warning = doc.add_table(rows=1, cols=1)
    set_table_geometry(warning, [9360])
    set_table_borders(warning, RED, 8)
    set_cell_shading(warning.cell(0, 0), RED_PALE)
    p = warning.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, "重要：", bold=True, color=RED)
    add_text(p, "降血压、心血管疾病预防、抗肿瘤、长期服用依赖、联合用药、服用周期等表述均已标为高风险，未填写批准终稿前不得生成正式旁白或机制动画。", color=RED)

    doc.add_page_break()
    add_heading(doc, "一、全局审核与授权清单", 1)
    global_rows = [
        ("药学审核", "K01–K18 批准终稿", "药师／合规", "☐ 未提交  ☐ 审核中  ☐ 已批准"),
        ("产品依据", "说明书／批准信息终稿", "业务／合规", "☐ 未提交  ☐ 已提交"),
        ("产品包装", "可可康高清正面／多面包装", "业务", "☐ 未授权  ☐ 已授权"),
        ("联合包装", "谷维素片、护肝片、转移因子包装", "业务", "☐ 未授权  ☐ 已授权"),
        ("卖点证据", "9.13g/100g、专利提取、GMP、先进技术", "业务／合规", "☐ 不使用  ☐ 待补  ☐ 已批准"),
        ("品牌资产", "大参林 Logo、内部培训声明", "品牌／业务", "☐ 待确认  ☐ 已批准"),
        ("产地／工厂", "大别山、中山可可康工厂／生产线素材", "业务", "☐ 不展示  ☐ 待补  ☐ 已授权"),
        ("服用周期", "1 个月／连续 2–3 个月口径", "药师／合规", "☐ 删除  ☐ 改写  ☐ 已批准"),
    ]
    gtable = doc.add_table(rows=1, cols=4)
    set_table_geometry(gtable, [1500, 3600, 1500, 2760])
    set_table_borders(gtable)
    headers = ["类别", "必须确认的内容", "责任方", "状态"]
    for i, h in enumerate(headers):
        set_cell_shading(gtable.cell(0, i), PALE)
        p = gtable.cell(0, i).paragraphs[0]
        style_paragraph(p, after=0)
        add_text(p, h, bold=True, color=GREEN_DARK)
    set_repeat_table_header(gtable.rows[0])
    for row in global_rows:
        cells = gtable.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=0)
            add_text(p, value, size=9.5, bold=(i == 0), color=GREEN_DARK if i == 0 else BLACK)
    set_table_geometry(gtable, [1500, 3600, 1500, 2760])
    set_table_borders(gtable)

    doc.add_page_break()
    add_heading(doc, "二、逐镜审核", 1)
    p = doc.add_paragraph()
    style_paragraph(p, after=10)
    add_text(p, "每镜旧课件口径仅用于定位。审核人应填写可直接生产的最终文字，不在此阶段调整画面版式或动画。", color=GRAY)

    for idx, scene in enumerate(SCENES):
        add_scene_block(doc, scene)
        if idx != len(SCENES) - 1:
            doc.add_page_break()

    doc.add_page_break()
    add_heading(doc, "三、最终签署与生产门禁", 1)
    sign_rows = [
        ("业务审核人", "", "部门／岗位", ""),
        ("药师／合规审核人", "", "审核日期", ""),
        ("批准脚本版本", "", "对应附件", ""),
        ("内容结论", "☐ 全部通过  ☐ 按批注修改后通过  ☐ 不通过", "是否允许制作金样", "☐ 否  ☐ 是，仅限内部验证"),
    ]
    stable = doc.add_table(rows=0, cols=4)
    for row in sign_rows:
        cells = stable.add_row().cells
        for i, value in enumerate(row):
            if i in (0, 2):
                set_cell_shading(cells[i], PALE)
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=0)
            add_text(p, value, bold=i in (0, 2), color=GREEN_DARK if i in (0, 2) else BLACK)
            if i in (1, 3) and not value:
                for _ in range(2):
                    cells[i].add_paragraph(" ")
    set_table_geometry(stable, [1700, 2980, 1700, 2980])
    set_table_borders(stable)

    add_heading(doc, "交接说明", 2)
    p = doc.add_paragraph()
    style_paragraph(p, after=7)
    add_text(p, "填写完成后，本表只作为绿色视频金样的内容输入。制作侧仍须使用授权资产、完整复刻成熟叙事框架，并提交整片签样；不得以本表通过替代视觉签样或第二商品批量验证。")

    final_note = doc.add_table(rows=1, cols=1)
    set_table_geometry(final_note, [9360])
    set_table_borders(final_note, GREEN, 8)
    set_cell_shading(final_note.cell(0, 0), PALE)
    p = final_note.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, "生产入口条件：", bold=True, color=GREEN_DARK)
    add_text(p, "K01–K18 全部有明确结论 + 高风险表述已有批准终稿 + 真包装与证据授权齐全 + 绿色完整视频金样方向已确认。")

    doc.core_properties.title = "可可康绿色视频金样逐镜内容审核确认表"
    doc.core_properties.subject = "K01–K18 业务与药师内容审核"
    doc.core_properties.author = "Chain Pharmacy Content Studio"
    doc.core_properties.keywords = "可可康, 绿色视频金样, 内容审核, K01-K18"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
