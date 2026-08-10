#!/usr/bin/env python3
"""Generate per-template business fill guides + product notebook filled sample.

Also refreshes settled filled-example mappings for product PPT templates so
business no longer sees 风热证 content as product format reference.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SETTLED = ROOT / "production-library/templates/settled"
OUT_SAMPLES = ROOT / "outputs/courseware-natural-import"

COMPONENT_BLANK_NAME = "构件化商品培训_业务提交_空白模板.docx"
COMPONENT_FILLED_NAME = "构件化商品培训_业务提交_填写参考.docx"

# compact_reference_guide preset: one quiet blue accent, compact tables, no decoration.
COMPONENT_FONT = "Arial Unicode MS"
COMPONENT_INK = RGBColor(0x25, 0x31, 0x3B)
COMPONENT_MUTED = RGBColor(0x67, 0x76, 0x82)
COMPONENT_ACCENT = RGBColor(0x2F, 0x5D, 0x7C)
COMPONENT_DEEP = RGBColor(0x1F, 0x46, 0x61)
COMPONENT_SOFT = "EAF3F8"
COMPONENT_SOFT_GRAY = "F5F7F9"
COMPONENT_LINE = "D9E2E8"

COMPONENT_CASES = (
    {
        "id": "A",
        "title": "资料核验闭环",
        "intent": "把三类已签样能力与一个新增页签组合成 7 页内部培训课件。",
        "pages": (
            ("1", "封面", "通用封面", "说明培训主题、对象与三项组合重点"),
            ("2", "为什么要先核验", "通用导语", "说明版本、来源与确认状态为何必须一致"),
            ("3", "核验对象总览", "绿色商品课型 · 商品总览能力", "列出对象编号、版本与使用范围"),
            ("4", "资料核验四步", "穿心莲课型 · 咨询框架能力", "按名称、版本、负责人、异常记录四步核验"),
            ("5", "交付证据阶梯", "速福达课型 · 证据阶梯能力", "依次核对包装文件、内容版本、视觉版本"),
            ("6", "异议与升级", "新增已登记页签", "回答临时替换、来源缺失时如何暂停与升级"),
            ("7", "核验闭环回顾", "通用总结", "收束识别、确认、留痕与升级"),
        ),
    },
    {
        "id": "B",
        "title": "陈列物料核验",
        "intent": "先看证据、再看物料，以非默认页序组成 6 页短课。",
        "pages": (
            ("1", "封面", "通用封面", "说明陈列物料核验任务"),
            ("2", "陈列物料状态检查", "通用数据导语", "点出版本、来源、状态三项检查"),
            ("3", "先看资料凭证", "速福达课型 · 证据阶梯能力", "核对原始文件、当前版本与取件状态"),
            ("4", "再看物料总览", "绿色商品课型 · 商品总览能力", "核对物料编号、版本与陈列范围"),
            ("5", "临时变更如何处理", "新增已登记页签", "说明换图、沿用旧版时的确认边界"),
            ("6", "陈列交付检查表", "通用总结", "汇总凭证、物料与变更状态"),
        ),
    },
    {
        "id": "C",
        "title": "交接与封存",
        "intent": "用两类已签样能力组成 5 页交接短课，不增加新页签。",
        "pages": (
            ("1", "封面", "通用封面", "说明交接与封存任务"),
            ("2", "为什么交接必须留痕", "通用导语", "说明同名文件不等于同一版本"),
            ("3", "交接四步检查", "穿心莲课型 · 咨询框架能力", "确认接收人、文件名、版本与回执"),
            ("4", "交接判定表", "通用总结", "区分可接收、需复核与需退回"),
            ("5", "最终封存凭证", "速福达课型 · 证据阶梯能力", "保存清单、文件哈希与验收结果"),
        ),
    },
)

# Recommended modules for business (not fixed page counts)
GUIDES: dict[str, dict] = {
    "kangaisen-lycopene-health-edu-v1": {
        "name_zh": "番茄红素成分健康科普 PPT（米白番茄红）",
        "outputs": "20 页可编辑 PPTX",
        "modules": [
            "封面与六章目录",
            "定义、来源与影响因素",
            "结构、比较数据与原理",
            "审核发现与证据",
            "建议流程、使用边界与注意事项",
            "应用场景、总结与结束页",
        ],
        "tips": [
            "这是 20 页成分健康科普课型，不是福尔商品课件4。",
            "业务按自然页面填写；WorkBuddy 负责转换为内部 107 个文字槽和 69 个图片绑定，业务不编辑 JSON。",
            "原康爱森正文和番茄红素参考图只用于看框架；正式新主题全部换文案、换图并绑定内容与视觉审批。",
            "涉及定义、机理、数据、功效、建议、风险或用法的内容必须来自业务审核稿，不得由 WorkBuddy 补写。",
        ],
        "chat_example": (
            "我要使用【番茄红素成分健康科普 PPT（米白番茄红）】。资料见 Word。\n"
            "请先整理 20 页内容初稿、待补字段和图片分工；在我确认内容与视觉之前，不要生成正式 PPTX。"
        ),
        "filled_source": "outputs/courseware-natural-import/成分健康科普_米白番茄红_业务提交_填写参考.docx",
        "blank_source": "outputs/courseware-natural-import/成分健康科普_米白番茄红_业务提交_空白模板.docx",
    },
    "health-video-reference-tech-v1": {
        "name_zh": "疾病科普视频（如风热证）",
        "outputs": "MP4 培训视频",
        "modules": [
            "开场",
            "基础认知",
            "病因与机理",
            "典型症状",
            "治疗思路",
            "用药与生活建议",
            "总结",
        ],
        "tips": [
            "**业务自助：** 在 WorkBuddy 对话里交完整 7 段审核稿；代理先出脚本/画面复核包。",
            "金样对照：`风热证_疾病科普视频_金样_v1.mp4`；正式换主题须补齐主题画面并完成当前载荷 SHA-256 审批。",
            "正式疾病科普片固定 7 段；缺段、内容缺口或待生成画面只交规划包，不冒充成片。",
            "旁白须药师/合规已审；正式成片用模板克隆药师声，禁止系统朗读。",
        ],
        "chat_example": (
            "我要用【疾病科普视频】模板，主题是【病名，如感冒】。\n"
            "内容围绕：开场、基础认知、病因机理、典型症状、调理建议、用药建议、总结…。\n"
            "请整理后直接生成培训视频（画面随主题换，不要只换声音）。"
        ),
        "agent_commands": [
            "# 先生成主题包，补齐 content_gaps / needs_generation 并在 review.html 全量过目",
            "python3 scripts/build_health_theme_package.py --theme <主题> --sections-json <path> --out-dir <theme-package目录>",
            "# approval.json 必须填写过目人、时间和当前 approved_payload_sha256",
            ".venv-qwen-tts/bin/python scripts/generate_business_video.py --template health --theme-package <theme-package目录> --with-tts --with-mp4 --copy-to-business-delivery",
            "# 仅规划（无 TTS 时）",
            "python3 scripts/generate_business_video.py --template health --mode plan --sections-json <path>",
        ],
        "filled_source": "outputs/video-training-natural-import/风热证健康知识视频培训_真实已填样本.docx",
        "blank_source": "outputs/video-training-natural-import/视频培训内容与素材提交_通用模板.docx",
    },
    "product-video-faithful-v1": {
        "name_zh": "商品培训视频（如辅酶 Q10）",
        "outputs": "MP4 培训视频",
        "modules": [
            "为什么要了解本商品",
            "商品基础信息",
            "核心功效 / 证据（审核稿）",
            "产品特点",
            "适宜人群",
            "联合用药（有几组写几组）",
            "总结",
        ],
        "tips": [
            "**业务自助：** 在 WorkBuddy 对话里说商品名+要点即可出片，不必找制作代跑。",
            "金样对照：`辅酶Q10_商品培训视频_金样_v1.mp4`；换商品走 full 分段重渲。",
            "一份内容一个商品；板块可删可重排。",
            "包装图必须用业务确认授权原图；无图只交规划包，禁止仿包装或正式渲染。正式旁白 = 审核原文 + 模板克隆声。",
        ],
        "chat_example": (
            "我要用【商品培训视频】模板，商品是【商品名】。\n"
            "内容围绕：核心功效…、产品特点…、适宜人群…、联合用药 2 组…。\n"
            "请整理后直接生成培训视频（画面随主题换，不要只换声音）。"
        ),
        "agent_commands": [
            "# 先出规划包与 product-approval.request.json（业务不碰命令）",
            "python3 scripts/generate_business_video.py --template product --mode plan --sections-json <path> --product-image <业务提供包装图>",
            "# 业务确认 8 段内容与包装授权后填写批准人/时间/授权凭证，再由 WorkBuddy 正式出片",
            ".venv-qwen-tts/bin/python scripts/generate_business_video.py --template product --sections-json <path> --with-tts --with-mp4 --product-image <业务确认授权包装图> --product-approval <已批准JSON> --copy-to-business-delivery",
        ],
        "filled_source": "outputs/video-training-natural-import/辅酶Q10商品培训视频_真实已填样本.docx",
        "blank_source": "outputs/video-training-natural-import/视频培训内容与素材提交_通用模板.docx",
    },
    "product-courseware-component-v1": {
        "name_zh": "灵活构件商品培训 PPT",
        "outputs": "可编辑 PPTX",
        "filled_source": f"outputs/courseware-natural-import/{COMPONENT_FILLED_NAME}",
        "blank_source": f"outputs/courseware-natural-import/{COMPONENT_BLANK_NAME}",
    },
    "product-courseware-green-v1": {
        "name_zh": "绿色单品 PPT（如金银花露）",
        "outputs": "可编辑 PPTX",
        "modules": [
            "商品介绍（成分/功能/用法）",
            "核心卖点（1～N 条）",
            "适宜人群（1～N 类）",
            "联合用药话术（1～N 组；2 组就 2 行）",
            "品种对标（可选，无则整节删）",
            "注意事项（可选）",
        ],
        "tips": [
            "用空白通用 Word，按上列板块起标题即可；不必写页码。",
            "联合用药只写真实有的组数，禁止为对齐示例空出第三行。",
            "先收「内容初稿 + 缺口清单」，确认后再出 PPTX。",
        ],
        "filled_source": "outputs/courseware-natural-import/商品培训_绿色单品_记事本式填写参考_两行联合用药.docx",
        "blank_source": "outputs/courseware-natural-import/培训课件内容与素材提交_通用模板.docx",
    },
    "disease-product-scenario-v1": {
        "name_zh": "疾病+商品场景 PPT（如穿心莲）",
        "outputs": "可编辑 PPTX",
        "modules": [
            "疾病/辨证知识",
            "商品知识",
            "销售场景与话术",
            "其他需要培训的板块（自定）",
        ],
        "tips": [
            "板块按资料完整度裁剪；没有的整节删除。",
            "填写参考若为他主题，只学结构不抄医学结论。",
        ],
        "filled_source": "outputs/courseware-natural-import/穿心莲内酯滴丸_商品培训课件_业务真实内容样本.docx",
        "blank_source": "outputs/courseware-natural-import/培训课件内容与素材提交_通用模板.docx",
    },
    "sufuda-mabaloshawei-product-courseware-3-v1": {
        "name_zh": "商品培训课件3（可编辑 PPT，速福达标准课型）",
        "outputs": "13 页可编辑 PPTX",
        "modules": [
            "课程开场 / 品类背景",
            "商品介绍与核心利益",
            "产品特点 / 证据",
            "适宜人群",
            "联合用药（有几组写几组）",
            "总结",
        ],
        "tips": [
            "Word 按 12 个主题内容单元写审核文案；适宜人群单元导出时拆为两页，因此终稿共 13 页。",
            "包装/Logo 必须使用业务授权原图；23 个非商品插图槽由 WorkBuddy 在内容确认后生成并逐槽绑定。",
            "当前新主题自助只交付 PPTX；MP4 仍在生产接入中，不承诺生成。",
        ],
        "filled_source": "outputs/courseware-natural-import/商品培训_绿色单品_记事本式填写参考_两行联合用药.docx",
        "blank_source": "outputs/courseware-natural-import/培训课件内容与素材提交_通用模板.docx",
    },
    "fuler-fanqiehongsu-product-courseware-4-v1": {
        "name_zh": "商品培训课件4（视频+PPT，番茄红素壳）",
        "outputs": "MP4 + 可编辑 PPTX",
        "modules": [
            "开场 / 商品介绍",
            "核心利益点",
            "原料与含量（若有审核稿）",
            "适宜人群",
            "关联用药（有几组写几组）",
            "总结",
        ],
        "tips": [
            "关联用药 note 在上、总结行标题完整句为课件4 语法，业务只需交审核文案。",
            "无包装图 → 槽位待补。",
        ],
        "filled_source": "outputs/courseware-natural-import/商品培训_绿色单品_记事本式填写参考_两行联合用药.docx",
        "blank_source": "outputs/courseware-natural-import/培训课件内容与素材提交_通用模板.docx",
    },
    "disease-health-shenke-blue-v1": {
        "name_zh": "疾病健康知识培训 PPT（参课蓝）",
        "outputs": "可编辑 PPTX",
        "modules": [
            "疾病概览（定义 + 病因）",
            "临床表现",
            "检查方法",
            "治疗用药（一般 / 全身 / 局部 / 对症表）",
            "用药注意事项（禁用须审核稿）",
            "专业关怀",
            "一页通（竖版，可选）",
        ],
        "tips": [
            "业务只写自然板块正文与授权图，不写页码/坐标。",
            "对症表可标重点药名；包装用授权原图，禁止伪造品牌包装。",
            "注意事项：仅「禁用」类表述在成品中标红；医学结论须业务/医学复核。",
            "先收「内容初稿 + 缺口清单」，确认后再出 PPTX。",
        ],
        "list_rules_note": "对症表/注意事项列表：有几条写几条；没有的板块整节删除。",
        "agent_commands": [
            "cd production-library/templates/settled/disease-health-shenke-blue-v1/generator",
            "npm install   # 首次或新 clone 后",
            "node build-editable.mjs content/<主题>.content.json",
        ],
        "filled_source": "outputs/courseware-natural-import/风热证培训课件内容与素材提交_真实已填样本.docx",
        "blank_source": "outputs/courseware-natural-import/培训课件内容与素材提交_通用模板.docx",
    },
}


def _set_run_gray(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _component_set_font(
    run,
    size: float,
    *,
    bold: bool = False,
    color: RGBColor = COMPONENT_INK,
) -> None:
    run.font.name = COMPONENT_FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("eastAsia", "ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{key}"), COMPONENT_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def _component_shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _component_set_cell_margins(
    cell,
    *,
    top: int = 80,
    start: int = 120,
    bottom: int = 80,
    end: int = 120,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _component_set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _component_set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _component_set_table_borders(table, color: str = COMPONENT_LINE) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "5")
        element.set(qn("w:color"), color)


def _component_configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = COMPONENT_FONT
    fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("eastAsia", "ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{key}"), COMPONENT_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = COMPONENT_INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, COMPONENT_ACCENT, 18, 10),
        ("Heading 2", 13, COMPONENT_DEEP, 14, 7),
        ("Heading 3", 12, COMPONENT_DEEP, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = COMPONENT_FONT
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("eastAsia", "ascii", "hAnsi", "cs"):
            fonts.set(qn(f"w:{key}"), COMPONENT_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _component_add_header_footer(doc: Document, label: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _component_set_font(header.add_run(label), 8.5, color=COMPONENT_MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _component_set_font(
        footer.add_run("业务提交用｜灵活构件商品培训 PPT"),
        8.5,
        color=COMPONENT_MUTED,
    )


def _component_add_callout(
    doc: Document,
    title: str,
    text: str,
    *,
    fill: str = COMPONENT_SOFT,
    add_spacer: bool = True,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    _component_set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    _component_shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    _component_set_font(p.add_run(f"{title}  "), 10.5, bold=True, color=COMPONENT_DEEP)
    _component_set_font(p.add_run(text), 10.5)
    if add_spacer:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)


def _component_fill_table_text(cell, text: str, *, bold: bool = False, size: float = 9.8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    _component_set_font(
        p.add_run(text),
        size,
        bold=bold,
        color=COMPONENT_DEEP if bold else COMPONENT_INK,
    )


def _component_add_cover(doc: Document, *, filled: bool) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(18)
    label.paragraph_format.space_after = Pt(4)
    _component_set_font(
        label.add_run("BUSINESS INPUT · FLEXIBLE COURSEWARE"),
        10,
        bold=True,
        color=COMPONENT_ACCENT,
    )
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    _component_set_font(
        title.add_run("灵活构件商品培训 PPT\n业务内容与素材提交"),
        25,
        bold=True,
    )
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _component_set_font(
        subtitle.add_run(
            "填写参考｜中性 A / B / C 多组合示范"
            if filled
            else "空白模板｜按目标、内容、素材自然填写"
        ),
        13,
        color=COMPONENT_MUTED,
    )
    table = doc.add_table(rows=4, cols=2)
    values = (
        ("商品 / 主题", "内部资料核验训练（中性示例）" if filled else "【请填写】"),
        ("交付目标", "让门店人员会识别版本、核对来源并保留确认记录" if filled else "【请填写：给谁用、希望学会什么】"),
        ("提交人 / 部门", "业务培训部 / 示例提交人" if filled else "【请填写】"),
        ("版本 / 日期", "结构参考 v1 / 2026-08-10" if filled else "【请填写】"),
    )
    for row, (left, right) in zip(table.rows, values):
        _component_shade_cell(row.cells[0], COMPONENT_SOFT)
        _component_fill_table_text(row.cells[0], left, bold=True, size=10.2)
        _component_fill_table_text(row.cells[1], right, size=10.2)
    _component_set_table_geometry(table, [2400, 6960])
    _component_set_table_borders(table)
    doc.add_paragraph()
    _component_add_callout(
        doc,
        "业务只需提供",
        "交付目标 + 已有业务内容 + 可授权素材。资料不完整可以先交，缺口明确标注即可。",
    )
    _component_add_callout(
        doc,
        "WorkBuddy 先做",
        "先给中文页签大纲、页签能力来源解释、待补清单、素材分工和一套统一视觉方案；业务确认前不生成正式 PPTX。",
        fill=COMPONENT_SOFT_GRAY,
    )
    doc.add_page_break()


def _component_add_workflow(doc: Document) -> None:
    doc.add_heading("业务怎么提交", level=1)
    table = doc.add_table(rows=3, cols=2)
    rows = (
        ("01 交付目标", "说明课件给谁、用于什么场景、希望培训后完成什么动作；页数可以不填。"),
        ("02 业务内容", "粘贴已有审核稿、数据、证据与边界；不完整处写“待补”，不要为了成稿自行补写医学内容。"),
        ("03 授权素材", "包装、Logo、真实人物、门店与证据截图由业务提供来源和授权；非商品插图由 WorkBuddy 按确定页签生成或绑定。"),
    )
    for row, (left, right) in zip(table.rows, rows):
        _component_shade_cell(row.cells[0], COMPONENT_SOFT)
        _component_fill_table_text(row.cells[0], left, bold=True, size=10)
        _component_fill_table_text(row.cells[1], right, size=10)
    _component_set_table_geometry(table, [2500, 6860])
    _component_set_table_borders(table)

    doc.add_heading("WorkBuddy 第一轮必须返回", level=2)
    first_review = (
        ("内容初稿与缺口", "忠实整理业务原文；所有待补、待确认和证据缺口单独列出。"),
        ("中文页签大纲", "根据内容选择页签与页序，用中文说明每一页准备讲什么。"),
        ("来源解释", "说明复用了哪类已签样信息结构；只复用能力，不复制来源文案、包装、原图或配色。"),
        ("单一视觉方案", "整份课件锁定一套视觉，不因页签来自不同课型而切换风格。"),
        ("素材分工", "逐项区分业务必须提供的授权素材与 WorkBuddy 可生成/绑定的非商品插图。"),
    )
    table = doc.add_table(rows=len(first_review), cols=2)
    for row, (left, right) in zip(table.rows, first_review):
        _component_shade_cell(row.cells[0], COMPONENT_SOFT_GRAY)
        _component_fill_table_text(row.cells[0], left, bold=True, size=9.7)
        _component_fill_table_text(row.cells[1], right, size=9.7)
    _component_set_table_geometry(table, [2350, 7010])
    _component_set_table_borders(table)
    doc.add_paragraph()
    _component_add_callout(
        doc,
        "确认顺序",
        "内容与中文页签 → 视觉与非商品插图绑定 → 正式商品图及授权 → 正式 PPTX → 逐页 QA。任何已确认载荷发生变化，都要重新确认。",
    )
    _component_add_callout(
        doc,
        "业务无需做",
        "无需选择“默认路线”，无需填写内部 JSON、页型 ID、坐标、卡片数量或生成命令。",
        fill=COMPONENT_SOFT_GRAY,
    )


def _component_add_input(doc: Document, *, filled: bool) -> None:
    doc.add_heading("完整业务输入示例" if filled else "业务填写区", level=1)
    goal_rows = (
        (
            "培训对象与场景",
            "连锁门店内部资料交接培训" if filled else "【请填写：谁会使用、在什么场景使用】",
        ),
        (
            "交付目标",
            "识别当前版本、核对来源、处理异常并形成可追溯记录"
            if filled
            else "【请填写：希望培训后能完成什么】",
        ),
        (
            "篇幅 / 节奏",
            "短课；页数由 WorkBuddy 根据内容建议" if filled else "【可不填；由 WorkBuddy 根据内容建议】",
        ),
        (
            "不能出现",
            "未确认结论、真实品牌信息和资料之外的推断"
            if filled
            else "【可不填：禁用说法、品牌边界、合规边界】",
        ),
    )
    table = doc.add_table(rows=len(goal_rows), cols=2)
    for row, (left, right) in zip(table.rows, goal_rows):
        _component_shade_cell(row.cells[0], COMPONENT_SOFT)
        _component_fill_table_text(row.cells[0], left, bold=True, size=9.9)
        _component_fill_table_text(row.cells[1], right, size=9.9)
    _component_set_table_geometry(table, [2500, 6860])
    _component_set_table_borders(table)

    doc.add_heading("业务内容（按自然逻辑写，不限栏目）", level=2)
    blocks = (
        (
            "对象与版本",
            "收到资料时，先核对对象名称、文件编号、版本号和使用范围。",
        ),
        (
            "核验规则",
            "依次确认原始文件、内容版本、视觉版本和批准记录是否一致。",
        ),
        (
            "异常升级",
            "来源找不到或版本冲突时，记录缺口并暂停使用；确认后再恢复。",
        ),
        (
            "交付留痕",
            "保存交接清单、批准人、确认时间、文件哈希和逐页验收结果。",
        ),
    ) if filled else (
        (
            "板块标题（自行命名，可复制或删除）",
            "【粘贴审核内容、数据、边界和证据；有几条写几条，缺失处写“待补”】",
        ),
        (
            "下一个板块标题（自行命名，可复制或删除）",
            "【继续按业务逻辑填写；不需要对应固定六栏目，也不需要写页码】",
        ),
    )
    for title, body in blocks:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph(body)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(5)
    _component_add_callout(
        doc,
        "内容边界",
        "WorkBuddy 可以整理、拆分和发现缺口，但不得自行编写医学功效、用法用量、联合推荐或业务未提供的结论。",
        fill=COMPONENT_SOFT_GRAY,
    )
    doc.add_page_break()

    doc.add_heading("素材与授权", level=1)
    asset_rows = (
        (
            "无品牌包装示意图",
            "业务提供",
            "内部 UAT 授权记录",
            "仅作结构参考；正式业务换为本商品授权包装图",
        ),
        (
            "版本确认记录",
            "业务提供",
            "业务资料库 / 确认人",
            "用于说明来源与版本",
        ),
        (
            "非商品插图",
            "WorkBuddy",
            "按已确认页签生成或绑定",
            "本示例无需外部非商品图片",
        ),
    ) if filled else (
        ("【素材名称】", "【业务提供 / WorkBuddy】", "【来源、授权人、范围】", "【关联哪段内容】"),
        ("【素材名称】", "【业务提供 / WorkBuddy】", "【来源、授权人、范围】", "【关联哪段内容】"),
        ("【素材名称】", "【业务提供 / WorkBuddy】", "【来源、授权人、范围】", "【关联哪段内容】"),
        ("【素材名称】", "【业务提供 / WorkBuddy】", "【来源、授权人、范围】", "【关联哪段内容】"),
    )
    table = doc.add_table(rows=1 + len(asset_rows), cols=4)
    headers = ("素材", "由谁提供", "来源 / 授权", "用途或说明")
    for index, text in enumerate(headers):
        _component_shade_cell(table.rows[0].cells[index], COMPONENT_SOFT)
        _component_fill_table_text(table.rows[0].cells[index], text, bold=True, size=9.3)
    for row, values in zip(table.rows[1:], asset_rows):
        for index, text in enumerate(values):
            _component_fill_table_text(row.cells[index], text, size=9.1)
    _component_set_table_geometry(table, [1850, 1750, 2750, 3010])
    _component_set_table_borders(table)

    doc.add_heading("复制给 WorkBuddy 的启动口令", level=2)
    prompt = (
        "我要做一份商品培训 PPT。交付目标、业务内容和素材见这份 Word。请根据内容选择合适的中文页签与页序，先给我内容初稿、待补字段、每个页签的能力来源解释、素材分工和一套统一视觉方案。在我确认内容、视觉和正式商品图授权前，不要生成正式 PPTX。我不填写内部 JSON 或页型 ID。"
    )
    _component_add_callout(doc, "启动口令", prompt)
    checklist = (
        "□ 目标、对象和使用场景已说明",
        "□ 所有未确认内容已标记待补",
        "□ 正式包装 / Logo / 人物 / 证据素材均有来源与授权",
        "□ 已要求先看中文页签大纲、来源解释和单一视觉方案",
    )
    for item in checklist:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)


def _component_add_case(doc: Document, case: dict, *, final: bool = False) -> None:
    pages = case["pages"]
    label = doc.add_paragraph()
    label.paragraph_format.space_after = Pt(3)
    _component_set_font(
        label.add_run(f"CASE {case['id']} · {len(pages)} 页组合"),
        10,
        bold=True,
        color=COMPONENT_ACCENT,
    )
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _component_set_font(title.add_run(case["title"]), 21, bold=True)
    intro = doc.add_paragraph(case["intent"])
    intro.paragraph_format.space_after = Pt(8)
    for run in intro.runs:
        _component_set_font(run, 10.5, color=COMPONENT_MUTED)

    table = doc.add_table(rows=1 + len(pages), cols=4)
    for index, text in enumerate(("页", "中文页签", "能力来源", "本页业务内容")):
        _component_shade_cell(table.rows[0].cells[index], COMPONENT_SOFT)
        _component_fill_table_text(table.rows[0].cells[index], text, bold=True, size=9.2)
    for row, values in zip(table.rows[1:], pages):
        for index, text in enumerate(values):
            _component_fill_table_text(row.cells[index], text, size=8.8)
    _component_set_table_geometry(table, [620, 2050, 2790, 3900])
    _component_set_table_borders(table)
    doc.add_paragraph()
    _component_add_callout(
        doc,
        "统一视觉",
        "整份只锁一套浅蓝商品培训视觉。来源只解释页签能力血缘，不带入来源课件的文案、包装、原图、母版或配色。",
        add_spacer=not final,
    )


def build_component_business_document(path: Path, *, filled: bool) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    _component_configure_styles(doc)
    _component_add_header_footer(doc, "填写参考" if filled else "空白模板")
    _component_add_cover(doc, filled=filled)
    _component_add_workflow(doc)
    doc.add_page_break()
    _component_add_input(doc, filled=filled)
    if filled:
        for index, case in enumerate(COMPONENT_CASES):
            doc.add_page_break()
            _component_add_case(doc, case, final=index == len(COMPONENT_CASES) - 1)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_component_business_words(*, sync: bool = False) -> tuple[Path, Path]:
    blank = OUT_SAMPLES / COMPONENT_BLANK_NAME
    filled = OUT_SAMPLES / COMPONENT_FILLED_NAME
    build_component_business_document(blank, filled=False)
    build_component_business_document(filled, filled=True)
    if sync:
        import shutil

        dest = SETTLED / "product-courseware-component-v1"
        shutil.copy2(blank, dest / "业务提交_空白模板.docx")
        shutil.copy2(filled, dest / "业务提交_填写参考.docx")
        write_guide("product-courseware-component-v1", GUIDES["product-courseware-component-v1"])
    return blank, filled


def build_product_notebook_filled_sample() -> Path:
    """Notebook-style product sample with exactly 2 combination rows."""
    out = OUT_SAMPLES / "商品培训_绿色单品_记事本式填写参考_两行联合用药.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("培训课件内容与素材提交")
    p = doc.add_paragraph()
    _set_run_gray(
        p,
        "填写参考｜只示范格式。示例医学/价格表述为演示占位，正式主题必须替换为公司审核稿与授权包装。",
    )
    p = doc.add_paragraph()
    _set_run_gray(
        p,
        "联合用药本参考只写 2 组 → 成片必须只出 2 行，禁止空白第三行。",
    )
    doc.add_paragraph("课件主题：示例商品A（演示用，非正式品规）")
    doc.add_paragraph("培训对象（可不填）：连锁药店门店员工")
    doc.add_paragraph("培训目标（可不填）：掌握审核稿中的商品信息与联合推荐话术边界")

    doc.add_heading("商品介绍", level=1)
    doc.add_paragraph("主要成分：【待替换为审核原文】")
    doc.add_paragraph("功能主治/定位：【待替换为审核原文】")
    doc.add_paragraph("用法用量：【待替换为审核原文】")
    doc.add_paragraph("【在此处直接粘贴本品授权包装图；没有可删除本行】")
    doc.add_paragraph("图片说明／来源（可不填）：公司授权包装")

    doc.add_heading("核心卖点", level=1)
    doc.add_paragraph("卖点一：【审核表述】")
    doc.add_paragraph("卖点二：【审核表述】")

    doc.add_heading("适宜人群", level=1)
    doc.add_paragraph("人群一：【审核表述】")
    doc.add_paragraph("人群二：【审核表述】")

    doc.add_heading("联合用药话术", level=1)
    doc.add_paragraph(
        "方案 1｜场景：【场景甲】｜组合：示例商品A + 搭档甲｜话术：【审核话术甲】"
    )
    doc.add_paragraph(
        "方案 2｜场景：【场景乙】｜组合：示例商品A + 搭档乙｜话术：【审核话术乙】"
    )
    doc.add_paragraph("【可粘贴搭档包装图；无图删除本行】")

    doc.add_heading("注意事项", level=1)
    doc.add_paragraph("【说明书或审核注意事项；没有整节可删】")

    doc.save(out)
    return out


def _write_component_guide(slug: str, meta: dict) -> Path:
    lines = [
        f"# 本课型怎么填 · {meta['name_zh']}",
        "",
        f"- 产物：{meta['outputs']}",
        "- 空白 Word：同目录 `业务提交_空白模板.docx`",
        "- 填写参考：同目录 `业务提交_填写参考.docx`（A / B / C 为中性组合示范，只学提交方法）",
        "",
        "## 业务只提供三类信息",
        "",
        "1. **交付目标：** 给谁培训、用于什么场景、希望培训后完成什么动作。",
        "2. **业务内容：** 粘贴已有审核稿、数据、证据与边界；资料不完整可直接标记“待补”。",
        "3. **授权素材：** 包装、Logo、真实人物、门店与证据截图由业务提供来源和授权；非商品插图由 WorkBuddy 负责生成或绑定。",
        "",
        "业务无需选择“默认路线”，无需填写内部 JSON、页型 ID、坐标、卡片数量或生成命令。也不要为了填满版式自行补写医学功效、用法用量或联合推荐。",
        "",
        "## WorkBuddy 第一轮必须先给业务确认",
        "",
        "1. 内容初稿与全部待补 / 待确认字段。",
        "2. 根据内容选择的**中文页签大纲与页序**。",
        "3. 每个页签的**能力来源解释**：只说明复用了哪类已签样信息结构，不复制来源文案、包装、原图或配色。",
        "4. 整份课件唯一的**单一视觉方案**；同一课件只锁一个 style pack。",
        "5. 素材分工：哪些必须由业务提供，哪些由 WorkBuddy 生成或绑定。",
        "",
        "## A / B / C 中性组合参考",
        "",
        "| 参考 | 培训任务 | 中文页签顺序 | 能力来源 |",
        "|---|---|---|---|",
        "| A · 7 页 | 资料核验闭环 | 封面 → 为什么先核验 → 对象总览 → 核验四步 → 证据阶梯 → 异议与升级 → 回顾 | 绿色商品总览 + 穿心莲咨询框架 + 速福达证据阶梯 + 新增异议页签 |",
        "| B · 6 页 | 陈列物料核验 | 封面 → 状态检查 → 先看凭证 → 再看总览 → 临时变更 → 检查表 | 速福达证据阶梯 + 绿色商品总览 + 新增变更页签 |",
        "| C · 5 页 | 交接与封存 | 封面 → 为什么留痕 → 交接四步 → 判定表 → 封存凭证 | 穿心莲咨询框架 + 速福达证据阶梯 |",
        "",
        "这三种不是让业务挑内部模板，而是说明 WorkBuddy 会按交付目标和业务内容组合不同页签、页序与页数。整份课件仍使用同一套视觉。",
        "",
        "## 确认与交付",
        "",
        "`内容与中文页签确认 → 视觉与非商品插图绑定确认 → 正式商品图及授权确认 → 正式 PPTX → 逐页 QA`",
        "",
        "任何已确认载荷发生变化，都必须重新确认；三道确认未齐全时不得发布正式 PPTX。",
        "",
        "## 对话示例（复制给 WorkBuddy）",
        "",
        "```",
        "我要做一份商品培训 PPT。交付目标、业务内容和素材见这份 Word。",
        "请根据内容选择合适的中文页签与页序，先给我内容初稿、待补字段、每个页签的能力来源解释、素材分工和一套统一视觉方案。",
        "在我确认内容、视觉和正式商品图授权前，不要生成正式 PPTX。我不填写内部 JSON 或页型 ID。",
        "```",
        "",
    ]
    path = SETTLED / slug / "本课型怎么填.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def write_guide(slug: str, meta: dict) -> Path:
    if slug == "product-courseware-component-v1":
        return _write_component_guide(slug, meta)
    lines = [
        f"# 本课型怎么填 · {meta['name_zh']}",
        "",
        f"- 产物：{meta['outputs']}",
        f"- 空白 Word：同目录 `业务提交_空白模板.docx`",
        f"- 填写参考：`../../03_填写参考/{slug}/业务提交_填写参考.docx`（仅学格式）",
        "",
        "## 推荐板块（可删可增）",
        "",
    ]
    for i, m in enumerate(meta["modules"], 1):
        lines.append(f"{i}. {m}")
    lines.extend(["", "## 填写要点", ""])
    for t in meta["tips"]:
        lines.append(f"- {t}")
    chat = meta.get("chat_example") or (
        f"我要用【{meta['name_zh']}】，主题是【病名或商品名】。\n"
        "请按金样整理后生成成片。"
    )
    lines.extend(
        [
            "",
            "## 对话示例（复制给 WorkBuddy）",
            "",
            "```",
            chat,
            "```",
            "",
            "## 列表 / 模块硬规则",
            "",
        ]
    )
    if meta.get("list_rules_note"):
        lines.append(f"- {meta['list_rules_note']}")
    else:
        lines.extend(
            [
                "- 有几条写几条；联合用药 2 组 → 成品 2 行",
                "- 禁止空行凑满金样示例数",
                "- 没有的板块整节删除",
            ]
        )
    if meta.get("agent_commands"):
        lines.extend(["", "## 代理出片命令（业务无需操作）", "", "```bash"])
        lines.extend(meta["agent_commands"])
        lines.append("```")
    lines.append("")
    path = SETTLED / slug / "本课型怎么填.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def sync_words() -> None:
    for slug, meta in GUIDES.items():
        dest_dir = SETTLED / slug
        blank_src = ROOT / meta["blank_source"]
        filled_src = ROOT / meta["filled_source"]
        if not blank_src.is_file():
            raise FileNotFoundError(blank_src)
        if not filled_src.is_file():
            raise FileNotFoundError(filled_src)
        import shutil

        shutil.copy2(blank_src, dest_dir / "业务提交_空白模板.docx")
        shutil.copy2(filled_src, dest_dir / "业务提交_填写参考.docx")
        write_guide(slug, meta)
        print(f"OK words+guide {slug}")


def main() -> None:
    from build_ingredient_health_business_words import build_sources

    build_sources(sync=False)
    for component_word in build_component_business_words(sync=False):
        print(f"Wrote {component_word.relative_to(ROOT)}")
    sample = build_product_notebook_filled_sample()
    print(f"Wrote {sample.relative_to(ROOT)}")
    sync_words()
    # update sync_settled mapping file note via rewrite of MAPPINGS is done in this script as SSOT
    catalog_note = SETTLED / "business-word-sources.json"
    catalog_note.write_text(
        __import__("json").dumps(
            {
                "version": "1.0.0",
                "note": "Authoritative Word sources for settled templates; regenerated by build_business_word_guides.py",
                "templates": {
                    slug: {
                        "blank": meta["blank_source"],
                        "filled": meta["filled_source"],
                        "name_zh": meta["name_zh"],
                    }
                    for slug, meta in GUIDES.items()
                },
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    print(f"Wrote {catalog_note.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
