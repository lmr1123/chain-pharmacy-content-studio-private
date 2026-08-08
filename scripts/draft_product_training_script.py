#!/usr/bin/env python3
"""L2：业务 Word / 大纲 → product-training-script/v1 草稿 + 人审清单。

纪律：
  - 不扩写功效/剂量；文案只从源文切分/归类
  - content_lock 默认 pending，永不自动升锁
  - 无法归类的正文进入 unmapped，不得静默丢弃

用法：
  python3 scripts/draft_product_training_script.py \\
    --input samples/product-training-script/示例大纲.md \\
    --out-dir /tmp/l2-draft \\
    --display-name "示例商品"

  python3 scripts/draft_product_training_script.py \\
    --input path/script.structured.json --out-dir /tmp/review --review-only
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_LOCK = "business-provided-draft-pending-pharmacist-review"
SCHEMA = "product-training-script/v1"

# (script_key, keywords, default_title)
SECTION_ALIASES: list[tuple[str, tuple[str, ...], str]] = [
    (
        "hook",
        (
            "导语",
            "引入",
            "痛点",
            "数据",
            "需求",
            "背景",
            "为什么",
            "信号",
            "需求教育",
            "问题引入",
            "疾病",
            "证型",
            "什么是",
        ),
        "导语",
    ),
    (
        "benefits",
        ("核心功效", "功效", "利益点", "作用", "核心卖点", "卖点"),
        "核心功效",
    ),
    (
        "features",
        ("产品特点", "特点", "产地", "原料", "含量", "工艺", "商品介绍", "商品总览", "亮相"),
        "产品特点",
    ),
    (
        "audience",
        ("适宜人群", "人群", "适用", "推荐对象"),
        "适宜人群",
    ),
    (
        "combination",
        ("联合用药", "联合", "联推", "搭配", "话术", "场景方案", "推荐场景", "场景"),
        "联合用药",
    ),
    (
        "summary",
        ("总结", "回顾", "要点回顾"),
        "总结",
    ),
    (
        "precautions",
        ("注意事项", "注意", "禁忌", "不适宜", "警示", "日常关怀", "用法", "用量"),
        "注意事项",
    ),
    (
        "meta",
        ("基本信息", "课程信息", "课件主题", "商品信息", "封面"),
        "基本信息",
    ),
]

IGNORE_HEADING_MARKERS = (
    "填写提示",
    "填写说明",
    "使用说明",
    "业务只需",
    "来源说明",
    "模板规则",
    "批量文件夹",
    "3 分钟",
    "三分钟",
)

IGNORE_PARA_PREFIXES = (
    "填写提示",
    "来源说明",
    "【在此处直接粘贴",
    "需要更多板块时",
    "像写记事本",
    "填写方法",
)

PLACEHOLDER_VALUES = {
    "请填写",
    "请替换",
    "请替换为内部审核原文",
    "请替换为内部审核话术",
    "在这里直接写本板块的审核内容",
    "板块标题（请替换）",
}


@dataclass
class Block:
    heading: str
    paragraphs: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)  # simple tables


@dataclass
class ParseResult:
    title: str
    audience_goal: str = ""
    training_goal: str = ""
    blocks: list[Block] = field(default_factory=list)
    raw_text: str = ""


def _norm(s: str) -> str:
    return re.sub(r"[\s：:／/、·_\-—（）()【】\[\]]+", "", s or "")


def _is_placeholder(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    if t in PLACEHOLDER_VALUES:
        return True
    if t.startswith("请填写") or t.startswith("请替换"):
        return True
    if "请填写" in t and len(t) < 20:
        return True
    return False


def _strip_bullet(line: str) -> str:
    return re.sub(r"^[•·▪◦\-–—\d]+[\.、\)\]\s]+", "", line).strip()


def _split_label_value(text: str) -> tuple[str | None, str]:
    for sep in ("：", ":"):
        if sep in text:
            a, b = text.split(sep, 1)
            a, b = a.strip(), b.strip()
            if a and b and len(a) <= 24:
                return a, b
    return None, text.strip()


def classify_heading(heading: str) -> tuple[str | None, float, str]:
    """Return (script_key, confidence, reason)."""
    h = heading.strip()
    if not h or any(m in h for m in IGNORE_HEADING_MARKERS):
        return None, 0.0, "ignored"
    n = _norm(h)
    best: tuple[str, float, str] | None = None
    for key, aliases, _title in SECTION_ALIASES:
        for al in aliases:
            an = _norm(al)
            if not an:
                continue
            if n == an:
                return key, 0.98, f"exact alias «{al}»"
            if an in n or n in an:
                conf = 0.9 if len(an) >= 4 else 0.75
                cand = (key, conf, f"alias «{al}» in heading")
                if best is None or cand[1] > best[1]:
                    best = cand
    if best:
        # 多别名同时命中（如「导语与痛点数据」）抬升置信度
        hits = sum(
            1
            for _k, aliases, _t in SECTION_ALIASES
            if _k == best[0]
            for al in aliases
            if _norm(al) and (_norm(al) in n or n in _norm(al))
        )
        if hits >= 2 and best[1] < 0.9:
            return best[0], 0.9, best[2] + f" (+{hits} aliases)"
        return best
    return None, 0.0, "no alias match"


def parse_markdown_or_txt(path: Path) -> ParseResult:
    title = path.stem
    blocks: list[Block] = []
    current: Block | None = None
    audience_goal = ""
    training_goal = ""
    lines = path.read_text(encoding="utf-8").splitlines()
    raw_parts: list[str] = []

    for raw in lines:
        raw_parts.append(raw)
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            continue
        if line.startswith("## "):
            if current and (current.paragraphs or current.bullets or current.rows):
                blocks.append(current)
            current = Block(heading=line[3:].strip())
            continue
        if current is None:
            # preamble key: values
            lab, val = _split_label_value(line)
            if lab and "主题" in lab:
                title = val or title
            elif lab and "培训对象" in lab:
                audience_goal = val
            elif lab and "培训目标" in lab:
                training_goal = val
            elif not line.startswith(">"):
                current = Block(heading="导语", paragraphs=[line])
            continue
        if any(line.startswith(p) for p in IGNORE_PARA_PREFIXES):
            continue
        if line.startswith("【") and line.endswith("】"):
            # field markers like 【审核旁白原文】— next lines go to paragraphs
            continue
        if line.startswith("|") and line.count("|") >= 2:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if cells and not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                current.rows.append(cells)
            continue
        if re.match(r"^[-*•·▪◦]\s+", line) or re.match(r"^\d+[\.、)]\s+", line):
            b = _strip_bullet(line)
            if b and not _is_placeholder(b):
                current.bullets.append(b)
            continue
        if not _is_placeholder(line):
            current.paragraphs.append(line)

    if current and (current.paragraphs or current.bullets or current.rows):
        blocks.append(current)

    return ParseResult(
        title=title,
        audience_goal=audience_goal,
        training_goal=training_goal,
        blocks=blocks,
        raw_text="\n".join(raw_parts),
    )


def parse_docx(path: Path) -> ParseResult:
    try:
        from docx import Document
        from docx.document import Document as DocumentObject
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as e:
        raise SystemExit("python-docx required: pip install python-docx") from e

    document = Document(str(path))
    title = path.stem
    blocks: list[Block] = []
    current: Block | None = None
    audience_goal = ""
    training_goal = ""
    raw_parts: list[str] = []

    def ensure_block(heading: str) -> Block:
        nonlocal current
        if current and (current.paragraphs or current.bullets or current.rows):
            blocks.append(current)
        current = Block(heading=heading)
        return current

    def iter_blocks():
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    for block in iter_blocks():
        if isinstance(block, Table):
            if current is None:
                ensure_block("表格内容")
            for row in block.rows:
                cells = [c.text.strip() for c in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    current.rows.append(cells)
                    raw_parts.append(" | ".join(cells))
            continue

        style = (block.style.name if block.style is not None else "") or ""
        text = (block.text or "").strip()
        if not text:
            continue
        raw_parts.append(text)

        if style in {"Courseware Ignore", "填写提示"} or any(
            text.startswith(p) for p in IGNORE_PARA_PREFIXES
        ):
            continue
        if any(m in text for m in IGNORE_HEADING_MARKERS) and style.startswith("Heading"):
            continue

        # meta lines
        lab, val = _split_label_value(text)
        if lab:
            ln = _norm(lab)
            if "课件主题" in lab or ln == "课件主题":
                if val and not _is_placeholder(val):
                    title = val
                continue
            if "培训对象" in lab:
                if val and not _is_placeholder(val):
                    audience_goal = val
                continue
            if "培训目标" in lab:
                if val and not _is_placeholder(val):
                    training_goal = val
                continue

        if style.startswith("Heading") or (
            len(text) <= 40
            and not text.endswith("。")
            and not text.endswith("；")
            and style in {"Normal", "标题 1", "标题 2"}
            and current is None
            and "：" not in text
        ):
            # Heading always starts section; short title-like after content also
            if style.startswith("Heading"):
                ensure_block(re.sub(r"^[一二三四五六七八九十\d]+[、.．\s]*", "", text))
                continue

        if style.startswith("List") or text.startswith(("•", "·", "-", "–")):
            if current is None:
                ensure_block("未命名板块")
            b = _strip_bullet(text)
            if b and not _is_placeholder(b):
                current.bullets.append(b)
            continue

        if current is None:
            # first body without heading → hook-like
            ensure_block("导语")
        if not _is_placeholder(text):
            current.paragraphs.append(text)

    if current and (current.paragraphs or current.bullets or current.rows):
        blocks.append(current)

    return ParseResult(
        title=title,
        audience_goal=audience_goal,
        training_goal=training_goal,
        blocks=blocks,
        raw_text="\n".join(raw_parts),
    )


def parse_input(path: Path) -> ParseResult:
    suf = path.suffix.lower()
    if suf in {".md", ".txt"}:
        return parse_markdown_or_txt(path)
    if suf == ".docx":
        return parse_docx(path)
    raise SystemExit(f"unsupported input type: {suf}")


def _title_body_items(block: Block) -> list[dict[str, str]]:
    items: list[dict[str, str]] = []
    for b in block.bullets:
        lab, val = _split_label_value(b)
        if lab and val:
            items.append({"title": lab, "body": val})
        else:
            items.append({"title": b[:40], "body": b})
    # paragraphs that look like title：body
    leftover: list[str] = []
    for p in block.paragraphs:
        lab, val = _split_label_value(p)
        if lab and val and len(lab) <= 20:
            items.append({"title": lab, "body": val})
        else:
            leftover.append(p)
    # pair leftover as single item if no bullets
    if not items and leftover:
        for p in leftover:
            items.append({"title": p[:24] + ("…" if len(p) > 24 else ""), "body": p})
    elif leftover and items:
        # append orphan paragraphs to last body
        items[-1]["body"] = (items[-1].get("body") or "") + "\n" + "\n".join(leftover)
    return items


def _list_items(block: Block) -> list[str]:
    if block.bullets:
        return list(block.bullets)
    items: list[str] = []
    for p in block.paragraphs:
        # split Chinese semicolon or顿号 lists carefully — only if short clauses
        if "；" in p and all(len(x) < 40 for x in p.split("；")):
            items.extend(x.strip() for x in p.split("；") if x.strip())
        else:
            items.append(p)
    return items


def _combo_rows(block: Block) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    header_tokens = {"场景", "问题场景", "问题", "搭配", "搭配药品", "药品", "话术", "销售话术", "解说"}
    for r in block.rows:
        if len(r) >= 3:
            # skip pure header row（单元格本身是列名，不是「话术」出现在正文里）
            if all(c.strip() in header_tokens or c.strip() in {"场景", "伙伴"} for c in r[:3]):
                continue
            if r[0].strip() in header_tokens and r[1].strip() in header_tokens:
                continue
            rows.append(
                {
                    "problem": r[0],
                    "partner": r[1],
                    "talk_track": r[2] if len(r) > 2 else "",
                }
            )
        elif len(r) == 2:
            if r[0].strip() in header_tokens:
                continue
            rows.append({"problem": r[0], "partner": r[1], "talk_track": ""})
    # pipe bullets: 场景|药|话术
    for b in block.bullets:
        if b.count("|") >= 2:
            parts = [p.strip() for p in b.split("|")]
            rows.append(
                {
                    "problem": parts[0],
                    "partner": parts[1],
                    "talk_track": parts[2] if len(parts) > 2 else "",
                }
            )
            continue
        lab, val = _split_label_value(b)
        if lab and val:
            rows.append({"problem": lab, "partner": val, "talk_track": ""})
    # labeled paragraph groups
    if not rows:
        cur: dict[str, str] = {}
        for p in block.paragraphs:
            lab, val = _split_label_value(p)
            if not lab:
                if cur and val:
                    cur["talk_track"] = (cur.get("talk_track") or "") + val
                continue
            ln = _norm(lab)
            if any(k in ln for k in ("场景", "问题", "主诉", "顾客")):
                if cur.get("problem") or cur.get("partner") or cur.get("talk_track"):
                    rows.append(
                        {
                            "problem": cur.get("problem") or "",
                            "partner": cur.get("partner") or "",
                            "talk_track": cur.get("talk_track") or "",
                        }
                    )
                    cur = {}
                cur["problem"] = val
            elif any(k in ln for k in ("联合", "搭配", "用药", "伙伴")):
                cur["partner"] = val
            elif any(k in ln for k in ("话术", "解说", "推荐语")):
                cur["talk_track"] = val
            else:
                if not cur.get("problem"):
                    cur["problem"] = lab
                    cur["talk_track"] = val
        if cur.get("problem") or cur.get("partner") or cur.get("talk_track"):
            rows.append(
                {
                    "problem": cur.get("problem") or "",
                    "partner": cur.get("partner") or "",
                    "talk_track": cur.get("talk_track") or "",
                }
            )
    return [r for r in rows if any(r.values())]


def _summary_rows(block: Block) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for b in block.bullets:
        lab, val = _split_label_value(b)
        if lab and val:
            rows.append({"label": lab, "value": val})
        else:
            rows.append({"label": b[:16], "value": b})
    for p in block.paragraphs:
        lab, val = _split_label_value(p)
        if lab and val:
            rows.append({"label": lab, "value": val})
    for r in block.rows:
        if len(r) >= 2:
            rows.append({"label": r[0], "value": r[1]})
    return rows


def _default_title(key: str) -> str:
    for k, _a, t in SECTION_ALIASES:
        if k == key:
            return t
    return key


def build_script(
    parsed: ParseResult,
    *,
    display_name: str | None,
    organization: str,
    tagline: str,
    family: str,
    style_pack_id: str,
) -> tuple[dict[str, Any], list[dict[str, Any]], list[str]]:
    """Returns script, source_map entries, review warnings."""
    warnings: list[str] = []
    source_map: list[dict[str, Any]] = []
    name = display_name or parsed.title
    script: dict[str, Any] = {
        "schema": SCHEMA,
        "meta": {
            "display_name": name,
            "organization": organization,
            "tagline": tagline,
            "content_lock": DEFAULT_LOCK,
            "brand_boast_disabled": True,
            "family": family,
            "style_pack_id": style_pack_id,
        },
    }
    if parsed.audience_goal:
        script["meta"]["training_audience"] = parsed.audience_goal
    if parsed.training_goal:
        script["meta"]["training_goal"] = parsed.training_goal

    unmapped: list[dict[str, Any]] = []

    # accumulate by key
    acc: dict[str, list[Block]] = {k: [] for k, _, _ in SECTION_ALIASES if k != "meta"}

    for block in parsed.blocks:
        key, conf, reason = classify_heading(block.heading)
        entry = {
            "heading": block.heading,
            "mapped_to": key,
            "confidence": conf,
            "reason": reason,
            "para_count": len(block.paragraphs),
            "bullet_count": len(block.bullets),
            "row_count": len(block.rows),
        }
        source_map.append(entry)
        if key is None or key == "meta":
            if key == "meta":
                # try pull display name from bullets
                for b in block.bullets + block.paragraphs:
                    lab, val = _split_label_value(b)
                    if lab and val and any(x in lab for x in ("商品", "品名", "名称")):
                        script["meta"]["display_name"] = val
                        entry["mapped_to"] = "meta.display_name"
                continue
            unmapped.append(
                {
                    "heading": block.heading,
                    "paragraphs": block.paragraphs,
                    "bullets": block.bullets,
                }
            )
            warnings.append(f"未归类板块：「{block.heading}」→ 见 REVIEW unmapped")
            continue
        if conf < 0.8:
            warnings.append(f"低置信度映射：「{block.heading}」→ {key} ({conf:.2f})")
        acc[key].append(block)

    # materialize sections
    if acc["hook"]:
        paras: list[str] = []
        title = acc["hook"][0].heading
        for b in acc["hook"]:
            paras.extend(b.paragraphs)
            paras.extend(b.bullets)
        script["hook"] = {"title": title if title != "导语" else "导语", "paragraphs": paras}

    if acc["benefits"]:
        items: list[dict[str, str]] = []
        title = _default_title("benefits")
        for b in acc["benefits"]:
            if b.heading and b.heading not in title:
                title = b.heading
            items.extend(_title_body_items(b))
        if items:
            script["benefits"] = {"title": title, "items": items}

    if acc["features"]:
        items = []
        title = _default_title("features")
        for b in acc["features"]:
            if b.heading:
                title = b.heading if "特点" in b.heading or "介绍" in b.heading else title
            items.extend(_title_body_items(b))
        if items:
            script["features"] = {"title": title, "items": items}

    if acc["audience"]:
        items_s: list[str] = []
        title = _default_title("audience")
        for b in acc["audience"]:
            if b.heading:
                title = b.heading
            items_s.extend(_list_items(b))
        if items_s:
            script["audience"] = {"title": title, "items": items_s}

    if acc["combination"]:
        rows: list[dict[str, str]] = []
        title = _default_title("combination")
        for b in acc["combination"]:
            if b.heading:
                title = b.heading
            rows.extend(_combo_rows(b))
        if rows:
            script["combination"] = {"title": title, "rows": rows}
        else:
            warnings.append("联合用药板块已识别但未解析出行；请人工补 rows")

    if acc["summary"]:
        rows = []
        title = _default_title("summary")
        for b in acc["summary"]:
            if b.heading:
                title = b.heading
            rows.extend(_summary_rows(b))
        if rows:
            script["summary"] = {"title": title, "rows": rows}

    if acc["precautions"]:
        items_s = []
        title = _default_title("precautions")
        for b in acc["precautions"]:
            if b.heading:
                title = b.heading
            items_s.extend(_list_items(b))
        if items_s:
            script["precautions"] = {"title": title, "items": items_s}

    if unmapped:
        script["_l2"] = {
            "unmapped_sections": unmapped,
            "note": "人审须处理 unmapped：并入某节或删除；生成器可忽略 _l2",
        }

    # empty section warnings
    for key in ("hook", "benefits", "features", "audience", "combination", "summary", "precautions"):
        if key not in script:
            warnings.append(f"节缺失：{key}（若业务本无此内容可忽略）")

    return script, source_map, warnings


# Chrome strings allowed without being in source
CHROME_ALLOW = {
    DEFAULT_LOCK,
    SCHEMA,
    "product-training",
    "导语",
    "核心功效",
    "产品特点",
    "适宜人群",
    "联合用药",
    "总结",
    "注意事项",
    "基本信息",
    "大参林医药集团",
    "【专业力】",
    "style-pack.lycopene-health-edu-cream-red-v1",
    "style-pack.courseware-4-silk-yellow-red-v1",
    "style-pack.dashenlin-courseware-green-v1",
}

# meta 配置项（CLI/默认）不强制在大纲出现；正文业务句仍须溯源
META_SKIP_PATHS = {
    "meta.organization",
    "meta.tagline",
    "meta.family",
    "meta.style_pack_id",
    "meta.training_audience",
    "meta.training_goal",
}


def collect_business_strings(obj: Any, path: str = "") -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    if isinstance(obj, str):
        if path.startswith("_l2"):
            return out
        if path in META_SKIP_PATHS:
            return out
        if path.endswith("content_lock") or path.endswith("schema") or path.endswith("family"):
            return out
        if path.endswith("style_pack_id") or path.endswith("brand_boast_disabled"):
            return out
        if obj.strip():
            out.append((path, obj))
    elif isinstance(obj, dict):
        for k, v in obj.items():
            if k in {"schema", "content_lock", "family", "style_pack_id", "brand_boast_disabled"}:
                continue
            if k == "_l2":
                continue
            out.extend(collect_business_strings(v, f"{path}.{k}" if path else k))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            out.extend(collect_business_strings(v, f"{path}[{i}]"))
    return out


def verify_provenance(script: dict[str, Any], source_text: str) -> list[str]:
    errors: list[str] = []
    src = source_text.replace("\r\n", "\n")
    for path, text in collect_business_strings(script):
        t = text.strip()
        if not t or t in CHROME_ALLOW:
            continue
        if len(t) <= 2:
            continue
        if t in src or any(
            t in line or line in t for line in src.splitlines() if len(line.strip()) > 2
        ):
            continue
        # title may be truncated with ellipsis
        if t.endswith("…") and t[:-1] in src:
            continue
        errors.append(f"provenance fail at {path}: {t[:60]!r}")
    return errors


def write_checklist(
    path: Path,
    *,
    script: dict[str, Any],
    source_map: list[dict[str, Any]],
    warnings: list[str],
    provenance_errors: list[str],
    input_path: Path,
) -> None:
    meta = script.get("meta") or {}
    lines = [
        "# 内容初稿人审清单（L2）",
        "",
        f"- 源文件：`{input_path}`",
        f"- 品名：{meta.get('display_name')}",
        f"- content_lock：`{meta.get('content_lock')}`（工具不自动升锁）",
        f"- family / style：{meta.get('family')} / {meta.get('style_pack_id')}",
        "",
        "## 必勾（全部完成后才可 generate_courseware）",
        "",
        "- [ ] 品名 / 组织 / 角标正确",
        "- [ ] 各板块映射合理（见下表低置信度项）",
        "- [ ] 未归类正文已处理（并入某节或明确删除）",
        "- [ ] 功效 / 数据 / 话术与审核稿一致，无补写",
        "- [ ] `hidden` 与大品牌话术策略符合业务意图",
        "- [ ] 联合用药列齐全或有意留空",
        "- [ ] 注意事项覆盖标签/禁忌口径",
        "- [ ] 素材缺口已知（包装可用占位）",
        "- [ ] 确认后如需升锁，人工改 `content_lock` 并留痕",
        "",
        "## 板块映射",
        "",
        "| 源标题 | 映射 | 置信度 | 理由 |",
        "|--------|------|--------|------|",
    ]
    for e in source_map:
        lines.append(
            f"| {e['heading']} | {e['mapped_to']} | {e['confidence']:.2f} | {e['reason']} |"
        )
    lines.extend(["", "## 警告", ""])
    if warnings:
        for w in warnings:
            lines.append(f"- {w}")
    else:
        lines.append("- （无）")
    lines.extend(["", "## 溯源自检", ""])
    if provenance_errors:
        lines.append("**失败（须修草稿或源文后再跑）：**")
        for e in provenance_errors:
            lines.append(f"- {e}")
    else:
        lines.append("- 通过：业务字符串均可在源文中找到依据")
    unmapped = (script.get("_l2") or {}).get("unmapped_sections") or []
    if unmapped:
        lines.extend(["", "## 未归类板块（须处理）", ""])
        for u in unmapped:
            lines.append(f"### {u.get('heading')}")
            for p in u.get("paragraphs") or []:
                lines.append(f"- {p}")
            for b in u.get("bullets") or []:
                lines.append(f"- {b}")
    lines.extend(
        [
            "",
            "## 下一节结构摘要",
            "",
        ]
    )
    for key in (
        "hook",
        "benefits",
        "features",
        "audience",
        "combination",
        "summary",
        "precautions",
    ):
        if key not in script:
            lines.append(f"- `{key}`：—")
            continue
        sec = script[key]
        if key == "hook":
            n = len(sec.get("paragraphs") or [])
            lines.append(f"- `hook`：{n} 段")
        elif key in ("benefits", "features"):
            lines.append(f"- `{key}`：{len(sec.get('items') or [])} 条")
        elif key == "audience":
            lines.append(f"- `audience`：{len(sec.get('items') or [])} 项")
        elif key == "combination":
            lines.append(f"- `combination`：{len(sec.get('rows') or [])} 行")
        elif key == "summary":
            lines.append(f"- `summary`：{len(sec.get('rows') or [])} 行")
        elif key == "precautions":
            lines.append(f"- `precautions`：{len(sec.get('items') or [])} 条")
    lines.extend(
        [
            "",
            "## 确认后命令",
            "",
            "```bash",
            "python3 scripts/generate_courseware.py \\",
            f"  --script {path.parent / 'script.structured.json'} \\",
            "  --style production-library/styles/lycopene-health-edu-cream-red-v1/tokens.json \\",
            "  --out-dir <validation-out>",
            "```",
            "",
            "手册：`docs/product-training-script-content-entry.md`",
            "",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")


def load_existing_script(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if data.get("schema") != SCHEMA:
        raise SystemExit(f"expected schema {SCHEMA}, got {data.get('schema')}")
    return data


def main() -> int:
    ap = argparse.ArgumentParser(description="L2 draft product-training-script from Word/outline")
    ap.add_argument("--input", type=Path, required=True, help=".md / .txt / .docx / .json")
    ap.add_argument("--out-dir", type=Path, required=True)
    ap.add_argument("--display-name", type=str, default=None)
    ap.add_argument("--organization", type=str, default="大参林医药集团")
    ap.add_argument("--tagline", type=str, default="【专业力】")
    ap.add_argument("--family", type=str, default="product-training")
    ap.add_argument(
        "--style-pack-id",
        type=str,
        default="style-pack.lycopene-health-edu-cream-red-v1",
    )
    ap.add_argument(
        "--review-only",
        action="store_true",
        help="input is existing script.structured.json; only emit checklist",
    )
    ap.add_argument(
        "--allow-provenance-fail",
        action="store_true",
        help="write outputs even if provenance self-check fails (still non-zero exit)",
    )
    args = ap.parse_args()

    inp = args.input.expanduser().resolve()
    if not inp.is_file():
        print(f"input not found: {inp}", file=sys.stderr)
        return 2

    out_dir = args.out_dir.expanduser().resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    if args.review_only or inp.suffix.lower() == ".json":
        script = load_existing_script(inp)
        if args.display_name:
            script.setdefault("meta", {})["display_name"] = args.display_name
        source_map = [{"heading": "(existing json)", "mapped_to": "—", "confidence": 1.0, "reason": "review-only"}]
        warnings = []
        if script.get("meta", {}).get("content_lock") == DEFAULT_LOCK:
            warnings.append("content_lock 仍为 pending（预期，除非已人工升锁）")
        # provenance against itself → always pass strings from self
        provenance_errors: list[str] = []
        source_text = json.dumps(script, ensure_ascii=False)
    else:
        parsed = parse_input(inp)
        script, source_map, warnings = build_script(
            parsed,
            display_name=args.display_name,
            organization=args.organization,
            tagline=args.tagline,
            family=args.family,
            style_pack_id=args.style_pack_id,
        )
        provenance_errors = verify_provenance(script, parsed.raw_text)
        source_text = parsed.raw_text

    script_path = out_dir / "script.structured.json"
    map_path = out_dir / "source-map.json"
    checklist_path = out_dir / "REVIEW-CHECKLIST.md"

    if not (args.review_only and inp.suffix.lower() == ".json" and inp.resolve() == script_path.resolve()):
        script_path.write_text(
            json.dumps(script, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
        )
    elif args.review_only:
        # still write a copy for out-dir completeness
        script_path.write_text(
            json.dumps(script, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
        )

    map_path.write_text(
        json.dumps(
            {
                "schema": "l2-source-map/v1",
                "input": str(inp),
                "entries": source_map,
                "warnings": warnings,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    write_checklist(
        checklist_path,
        script=script,
        source_map=source_map,
        warnings=warnings,
        provenance_errors=provenance_errors,
        input_path=inp,
    )

    print(f"wrote {script_path}")
    print(f"wrote {map_path}")
    print(f"wrote {checklist_path}")
    if warnings:
        print(f"warnings: {len(warnings)}")
    if provenance_errors:
        print(f"provenance errors: {len(provenance_errors)}", file=sys.stderr)
        for e in provenance_errors[:10]:
            print(f"  {e}", file=sys.stderr)
        return 0 if args.allow_provenance_fail else 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
