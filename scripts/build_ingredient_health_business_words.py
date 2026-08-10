#!/usr/bin/env python3
"""Build the two business-facing Word inputs for the 20-page ingredient courseware.

The Word files are the business interface. WorkBuddy converts their natural-language
sections into the internal 107 text / 69 image bindings; business users never edit
the JSON shape-id contract.
"""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "outputs/courseware-natural-import"
SETTLED_DIR = (
    ROOT
    / "production-library/templates/settled/kangaisen-lycopene-health-edu-v1"
)
BLANK_NAME = "成分健康科普_米白番茄红_业务提交_空白模板.docx"
FILLED_NAME = "成分健康科普_米白番茄红_业务提交_填写参考.docx"

# compact_reference_guide preset, with one named Chinese-font/red-courseware override.
FONT = "Arial Unicode MS"
INK = RGBColor(0x2C, 0x2C, 0x2C)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
ACCENT = RGBColor(0xD3, 0x2F, 0x2F)
DEEP = RGBColor(0x9F, 0x22, 0x22)
SOFT_RED = "FCEBE8"
SOFT_CREAM = "FFF8F2"
LINE = "E8DDD6"


PAGE_SPECS = [
    (1, "封面", "cover", "课程主题、解释性副标题、培训对象、讲师/部门、1 张主题主视觉"),
    (2, "目录", "toc", "6 个章节名称；章节顺序可按审核资料调整；目录图标由 WorkBuddy 配置"),
    (3, "章节一扉页", "chapter", "章节编号、中文章名、英文短标题、1 张章节主题图"),
    (4, "核心定义", "two_card_media", "定义卡、特性卡、名词解释、证据来源、1 张结构或概念示意图"),
    (5, "来源 / 分布", "icon_grid", "3–6 个来源或分类：每项含名称、简述、图片需求与来源"),
    (6, "影响因素", "two_feature", "因素 A、因素 B：各含标题、说明、证据出处和配图需求"),
    (7, "章节二扉页", "chapter", "章节编号、中文章名、英文短标题、1 张章节主题图"),
    (8, "数据或特性比较", "compare_chart", "比较结论、指标名称、全部数值与单位、基准、出处、图表配色说明"),
    (9, "原理 / 机制", "mechanism_pair", "左侧前提、右侧过程、结果边界、证据出处、1 张机制示意图"),
    (10, "章节三扉页", "chapter", "章节编号、中文章名、英文短标题、1 张章节主题图"),
    (11, "研究或应用发现（一）", "list_media", "页面结论、3 个已审核要点、证据出处、1 张场景或数据图"),
    (12, "研究或应用发现（二）", "list_media", "页面结论、2–3 个已审核要点、证据出处、1 张场景或数据图"),
    (13, "章节四扉页", "chapter", "章节编号、中文章名、英文短标题、1 张章节主题图"),
    (14, "建议流程", "three_process", "3 个步骤：每步含动作、说明、适用边界、证据出处和图片需求"),
    (15, "使用边界与注意事项", "dose_notice", "审核建议、适用边界、注意事项、禁忌/风险、证据出处；不得自行编写剂量"),
    (16, "章节五扉页", "chapter", "章节编号、中文章名、英文短标题、1 张章节主题图"),
    (17, "应用场景", "app_grid", "4 个应用场景：每项含标题、说明、图片需求与来源/生成方式"),
    (18, "章节六扉页", "chapter", "章节编号、中文章名、英文短标题、1 张章节主题图"),
    (19, "总结", "summary_split", "左栏 2–3 个知识结论；右栏行动建议、边界提示和审核来源"),
    (20, "结束页", "end", "结束标题、英文短句、免责声明/联系信息、1 张收尾视觉"),
]


REFERENCE_VALUES = {
    1: "主题：示例成分 A｜副标题：从定义到应用的结构示范｜对象：门店培训人员｜主视觉：抽象成分图（生成）",
    2: "章节：认识成分、结构与原理、审核发现、使用边界、应用场景、总结；正式稿可删改章名，但需保持 20 页课型结构。",
    3: "01｜认识示例成分 A｜Definition & Source｜章节图：原料与结构的非品牌插画（生成）",
    4: "定义卡：示例成分 A 是虚构占位名；特性卡：本页仅示范信息层级。正式稿粘贴公司审核定义、名词解释和证据编号。",
    5: "来源分类一 / 二 / 三：分别填写真实来源、1 句说明和授权图；若用生成图，写清生成主题与禁用元素。",
    6: "因素 A / 因素 B：示范两类影响因素的对照写法；正式稿需给出审核结论、适用边界与证据出处。",
    7: "02｜结构与原理｜Structure & Principles｜章节图：抽象结构线稿（生成）",
    8: "指标 A / B / C：填写数值、单位、基准和出处；没有正式数据时，本页标记待补，不能用示例数字进入终稿。",
    9: "前提 → 过程 → 结果：只示范逻辑链。正式稿由医学/合规审核提供完整机制表述和引用。",
    10: "03｜审核发现｜Reviewed Findings｜章节图：资料与数据卡片场景（生成）",
    11: "结论句 + 要点 A / B / C + 证据编号；本参考不提供任何真实功效结论，正式内容以业务审核稿为准。",
    12: "第二组结论 + 要点 A / B；可与第 11 页形成两类应用或两组研究对照，必须分别标注来源。",
    13: "04｜使用边界｜Use & Boundaries｜章节图：步骤与提示标识（生成）",
    14: "步骤 1 识别场景｜步骤 2 核对资料｜步骤 3 按审核口径讲解；正式流程由业务确认。",
    15: "审核建议、适用边界、注意事项、禁忌/风险四栏均需业务提供；不得由 WorkBuddy补写剂量、功效或用法。",
    16: "05｜应用场景｜Applications｜章节图：四类场景拼图（生成）",
    17: "场景 A / B / C / D：每项 1 句用途说明 + 1 张图片；品牌、包装、证据截图必须由业务提供授权原图。",
    18: "06｜总结与展望｜Summary & Outlook｜章节图：收束与行动提示（生成）",
    19: "左栏：三个审核知识结论；右栏：一条行动建议 + 一条合规边界。引用编号与正文保持一致。",
    20: "结束标题：感谢学习｜免责声明：本课件仅用于内部培训，具体口径以公司审核资料为准｜收尾图：抽象主题视觉（生成）",
}


def _set_font(run, size: float, *, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_width(table, widths: list[int], *, indent: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DEEP, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_page_field(doc: Document, label: str, value: str, *, filled: bool) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    label_run = p.add_run(f"{label}：")
    _set_font(label_run, 10.5, bold=True, color=DEEP)
    body = value if filled else "【请填写；资料不完整可写“待补”，WorkBuddy 会列入缺口清单】"
    body_run = p.add_run(body)
    _set_font(body_run, 10.5, color=INK if filled else MUTED)


def _add_callout(doc: Document, title: str, text: str, *, fill: str = SOFT_CREAM) -> None:
    table = doc.add_table(rows=1, cols=1)
    cant_split = OxmlElement("w:cantSplit")
    table.rows[0]._tr.get_or_add_trPr().append(cant_split)
    _set_table_width(table, [9360])
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    title_run = p.add_run(f"{title}  ")
    _set_font(title_run, 10.5, bold=True, color=DEEP)
    text_run = p.add_run(text)
    _set_font(text_run, 10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_header_footer(doc: Document, label: str) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(label)
    _set_font(run, 8.5, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("业务提交用｜20 页米白番茄红成分健康科普课型")
    _set_font(run, 8.5, color=MUTED)


def _add_cover(doc: Document, *, filled: bool) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("BUSINESS INPUT · 成分健康科普")
    _set_font(r, 10, bold=True, color=ACCENT)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("米白番茄红 20 页课型\n业务内容与素材提交")
    _set_font(r, 25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("空白模板" if not filled else "填写参考｜仅示范结构，不含真实医学结论")
    _set_font(r, 13, color=MUTED)
    table = doc.add_table(rows=4, cols=2)
    values = [
        ("课件主题", "【请填写】" if not filled else "示例成分 A（虚构占位）"),
        ("提交人 / 部门", "【请填写】" if not filled else "业务培训部 / 示例提交人"),
        ("内容审核人", "【请填写】" if not filled else "待正式主题提交时填写"),
        ("版本 / 日期", "【请填写】" if not filled else "结构示范 v1 / 2026-08-09"),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        _shade_cell(row.cells[0], SOFT_RED)
        for index, cell in enumerate(row.cells):
            for run in cell.paragraphs[0].runs:
                _set_font(run, 10.5, bold=index == 0, color=DEEP if index == 0 else INK)
    _set_table_width(table, [2700, 6660])
    doc.add_paragraph()
    _add_callout(
        doc,
        "提交原则",
        "业务可先交残缺资料；WorkBuddy 先整理初稿与待补字段。正式 PPTX 只在 20 页文字、69 个图片绑定、内容确认和视觉确认全部完成后生成。",
    )
    _add_callout(
        doc,
        "内容边界",
        "不要复制康爱森金样正文或原图；不要自行编写功效、剂量、用法或医学结论。品牌、证据、人物、商品与实景图须由业务提供授权来源。",
        fill=SOFT_RED,
    )
    doc.add_page_break()


def _add_workflow(doc: Document) -> None:
    doc.add_heading("业务怎么提交", level=1)
    for number, (title, text) in enumerate(
        (
            ("选课型", "在门户选择“番茄红素成分健康科普 PPT（米白番茄红）”。"),
            ("交内容", "按下面 20 页自然填写；不必数 107 个内部文字框，也不必编辑 JSON。"),
            ("补素材", "业务提供必须真实/授权的图；非品牌插图由 WorkBuddy 按槽生成并给代表图复核。"),
            ("两道确认", "先确认内容，再确认视觉。任一确认后发生改动都需重新确认。"),
            ("正式生成", "WorkBuddy 导出 20 页可编辑 PPTX，并完成逐页布局、残留和素材哈希 QA。"),
        ),
        1,
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{number}. {title}  ")
        _set_font(r, 11, bold=True, color=DEEP)
        r = p.add_run(text)
        _set_font(r, 11, color=INK)
    _add_callout(
        doc,
        "图片分工",
        "必须由业务提供：Logo、品牌/包装、证据截图、真实人物或门店等不可伪造素材。可由 WorkBuddy 生成：通用成分示意、抽象机制、装饰图标和非品牌场景插画；生成记录也要进入授权清单。",
    )


def _add_page_specs(doc: Document, *, filled: bool) -> None:
    doc.add_heading("20 页内容与素材", level=1)
    doc.add_paragraph(
        "每页先写审核内容，再写图片来源或生成需求。没有资料时保留“待补”并说明负责人；不要为了填满版式编写内容。"
    )
    chapter_starts = {3, 7, 10, 18} if filled else {3, 7, 10, 13, 16, 18}
    for slide, title, page_type, requirement in PAGE_SPECS:
        if slide in chapter_starts and slide != 3:
            doc.add_page_break()
        doc.add_heading(f"第 {slide:02d} 页｜{title}", level=2)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"固定页型：{page_type}")
        _set_font(r, 9.5, bold=True, color=MUTED)
        _add_page_field(doc, "本页需提交", requirement, filled=True)
        _add_page_field(
            doc,
            "审核正文",
            REFERENCE_VALUES[slide],
            filled=filled,
        )
        image_value = (
            "已标出由业务提供的授权图与可由 WorkBuddy 生成的非品牌插图；正式绑定以视觉复核清单为准。"
            if filled
            else "【逐项写：图片名称｜业务提供/WorkBuddy 生成｜来源或生成依据｜授权人】"
        )
        _add_page_field(doc, "图片与来源", image_value, filled=True)
        evidence_value = (
            "示例仅写结构；正式主题须填写公司审核稿编号、说明书/文献/内部资料路径与审核人。"
            if filled
            else "【填写证据编号、文件名/链接、页码或内部资料路径；无证据写待补】"
        )
        _add_page_field(doc, "证据 / 审核", evidence_value, filled=True)
        if slide in {4, 11, 14, 19}:
            _add_callout(
                doc,
                "医学与合规提醒",
                "本页可能涉及定义、机理、数据、功效、建议或风险；只能使用已审核原文，不得由 WorkBuddy 自行补写。",
                fill=SOFT_RED,
            )


def _add_final_checks(doc: Document, *, filled: bool) -> None:
    if not filled:
        doc.add_page_break()
    doc.add_heading("提交前复核", level=1)
    checks = (
        "20 页主题、标题和正文都已填写，或明确标注待补负责人",
        "所有数据、比较、机理、建议和注意事项都有审核来源",
        "业务提供的图片均写明授权人、来源和使用范围",
        "没有把康爱森金样文案、图片、品牌或示例数字带入新主题",
        "没有要求 WorkBuddy 伪造包装、Logo、人物、证据截图或医学结论",
        "已确认内容初稿；已查看代表性插图和关键页预览",
    )
    for index, text in enumerate(checks, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"检查 {index}  ")
        _set_font(r, 10.5, bold=True, color=DEEP)
        r = p.add_run(text)
        _set_font(r, 10.5, color=INK)
        r = p.add_run("    □ 已确认")
        _set_font(r, 10.5, bold=True, color=ACCENT)
    _add_callout(
        doc,
        "给 WorkBuddy 的启动口令",
        "我要使用【番茄红素成分健康科普 PPT（米白番茄红）】。资料见这份 Word。请先整理 20 页内容初稿、待补字段和图片分工；在我确认内容与视觉之前，不要生成正式 PPTX。",
    )


def build_document(path: Path, *, filled: bool) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    _configure_styles(doc)
    _add_header_footer(doc, "填写参考" if filled else "空白模板")
    _add_cover(doc, filled=filled)
    _add_workflow(doc)
    _add_page_specs(doc, filled=filled)
    _add_final_checks(doc, filled=filled)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_sources(*, sync: bool = False) -> tuple[Path, Path]:
    blank = SOURCE_DIR / BLANK_NAME
    filled = SOURCE_DIR / FILLED_NAME
    build_document(blank, filled=False)
    build_document(filled, filled=True)
    if sync:
        SETTLED_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(blank, SETTLED_DIR / "业务提交_空白模板.docx")
        shutil.copy2(filled, SETTLED_DIR / "业务提交_填写参考.docx")
    return blank, filled


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--sync", action="store_true", help="also refresh settled copies")
    args = parser.parse_args()
    for path in build_sources(sync=args.sync):
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
