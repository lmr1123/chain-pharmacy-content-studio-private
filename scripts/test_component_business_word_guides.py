#!/usr/bin/env python3
"""Truth and layout contract for component-route business Word guides."""

from __future__ import annotations

import json
import sys
import tempfile
import unittest
from xml.etree import ElementTree as ET
from pathlib import Path
from zipfile import ZipFile, is_zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import build_business_word_guides as builder  # noqa: E402


SLUG = "product-courseware-component-v1"
SETTLED = ROOT / "production-library/templates/settled" / SLUG
BLANK_SOURCE = ROOT / "outputs/courseware-natural-import" / builder.COMPONENT_BLANK_NAME
FILLED_SOURCE = ROOT / "outputs/courseware-natural-import" / builder.COMPONENT_FILLED_NAME


def _docx_text(path: Path) -> str:
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    parts: list[str] = []
    start = 0
    while True:
        left = xml.find("<w:t", start)
        if left < 0:
            break
        left = xml.find(">", left) + 1
        right = xml.find("</w:t>", left)
        parts.append(xml[left:right])
        start = right + 6
    return "\n".join(parts)


def _registry_entry() -> dict:
    data = json.loads((ROOT / "production-library/registries/templates.json").read_text())
    return next(item for item in data["items"] if item["id"] == f"template.{SLUG}")


class ComponentBusinessWordGuidesTest(unittest.TestCase):
    def test_generated_words_are_neutral_and_business_facing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "blank.docx"
            filled = tmp_path / "filled.docx"
            builder.build_component_business_document(blank, filled=False)
            builder.build_component_business_document(filled, filled=True)

            blank_text = _docx_text(blank)
            filled_text = _docx_text(filled)
            for text in (blank_text, filled_text):
                for required in (
                    "交付目标",
                    "业务内容",
                    "授权素材",
                    "中文页签大纲",
                    "来源解释",
                    "单一视觉方案",
                    "无需选择“默认路线”",
                    "内部 JSON",
                    "页型 ID",
                ):
                    self.assertIn(required, text)
                for forbidden in (
                    "金银花露",
                    "绿色单品",
                    "示例商品A",
                    "联合用药本参考",
                    "构件化商品培训 PPT（默认主路径）",
                ):
                    self.assertNotIn(forbidden, text)

            for required in (
                "CASE A · 7 页组合",
                "CASE B · 6 页组合",
                "CASE C · 5 页组合",
                "绿色商品课型 · 商品总览能力",
                "穿心莲课型 · 咨询框架能力",
                "速福达课型 · 证据阶梯能力",
                "新增已登记页签",
            ):
                self.assertIn(required, filled_text)

    def test_compact_reference_guide_geometry_is_exact(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "layout.docx"
            builder.build_component_business_document(path, filled=True)
            doc = Document(path)
            section = doc.sections[0]
            self.assertEqual(section.page_width, Inches(8.5))
            self.assertEqual(section.page_height, Inches(11))
            for value in (
                section.top_margin,
                section.right_margin,
                section.bottom_margin,
                section.left_margin,
            ):
                self.assertEqual(value, Inches(1))
            # OOXML stores these distances as integer twips; 0.492 in rounds to 708 twips.
            self.assertEqual(section.header_distance, Twips(708))
            self.assertEqual(section.footer_distance, Twips(708))

            normal = doc.styles["Normal"]
            self.assertEqual(normal.font.name, builder.COMPONENT_FONT)
            self.assertEqual(normal.font.size, Pt(11))
            self.assertEqual(normal.paragraph_format.space_after, Pt(6))
            self.assertEqual(normal.paragraph_format.line_spacing, 1.25)
            expected_headings = {
                "Heading 1": (Pt(16), Pt(18), Pt(10)),
                "Heading 2": (Pt(13), Pt(14), Pt(7)),
                "Heading 3": (Pt(12), Pt(10), Pt(5)),
            }
            for name, values in expected_headings.items():
                style = doc.styles[name]
                self.assertEqual(style.font.name, builder.COMPONENT_FONT)
                self.assertEqual(style.font.size, values[0])
                self.assertEqual(style.paragraph_format.space_before, values[1])
                self.assertEqual(style.paragraph_format.space_after, values[2])

            for table in doc.tables:
                tbl_pr = table._tbl.tblPr
                self.assertEqual(tbl_pr.find(qn("w:tblW")).get(qn("w:w")), "9360")
                self.assertEqual(tbl_pr.find(qn("w:tblInd")).get(qn("w:w")), "120")
                widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
                self.assertEqual(sum(widths), 9360)
                for cell in table.rows[0].cells:
                    margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
                    self.assertEqual(margins.find(qn("w:top")).get(qn("w:w")), "80")
                    self.assertEqual(margins.find(qn("w:bottom")).get(qn("w:w")), "80")
                    self.assertEqual(margins.find(qn("w:start")).get(qn("w:w")), "120")
                    self.assertEqual(margins.find(qn("w:end")).get(qn("w:w")), "120")

    def test_authoritative_sources_manifest_registry_and_settled_match(self) -> None:
        expected_blank = f"outputs/courseware-natural-import/{builder.COMPONENT_BLANK_NAME}"
        expected_filled = f"outputs/courseware-natural-import/{builder.COMPONENT_FILLED_NAME}"
        manifest = json.loads((SETTLED / "manifest.json").read_text())
        business_input = manifest["business_input"]
        self.assertEqual(business_input["blank_source"], expected_blank)
        self.assertEqual(business_input["example_source"], expected_filled)
        self.assertEqual(
            business_input["business_provides"],
            ["交付目标", "业务内容", "授权素材"],
        )
        self.assertIn("中文页签大纲与页序", business_input["workbuddy_first_review"])
        self.assertIn("页签能力来源解释", business_input["workbuddy_first_review"])
        self.assertEqual(business_input["style_rule"], "one-style-pack-per-courseware")

        source_map = json.loads(
            (ROOT / "production-library/templates/settled/business-word-sources.json").read_text()
        )["templates"][SLUG]
        self.assertEqual(source_map["blank"], expected_blank)
        self.assertEqual(source_map["filled"], expected_filled)

        registry_sources = _registry_entry()["business_input_sources"]
        self.assertEqual(registry_sources["blank_word"], expected_blank)
        self.assertEqual(registry_sources["filled_example"], expected_filled)

        for source, copy in (
            (BLANK_SOURCE, SETTLED / "业务提交_空白模板.docx"),
            (FILLED_SOURCE, SETTLED / "业务提交_填写参考.docx"),
        ):
            self.assertTrue(source.is_file())
            self.assertTrue(copy.is_file())
            self.assertEqual(source.read_bytes(), copy.read_bytes())
            self.assertTrue(is_zipfile(source))
            with ZipFile(source) as archive:
                self.assertIsNone(archive.testzip())

    def test_static_pagination_keep_next_and_empty_paragraphs(self) -> None:
        namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns = {"w": namespace}
        for path, expected_pages in (
            (SETTLED / "业务提交_空白模板.docx", 4),
            (SETTLED / "业务提交_填写参考.docx", 7),
        ):
            with ZipFile(path) as archive:
                root = ET.fromstring(archive.read("word/document.xml"))
            body = root.find("w:body", ns)
            self.assertIsNotNone(body)
            page_breaks = root.findall(".//w:br[@w:type='page']", ns)
            self.assertEqual(len(page_breaks) + 1, expected_pages)

            consecutive_empty = 0
            max_consecutive_empty = 0
            for child in body:
                if child.tag != qn("w:p"):
                    consecutive_empty = 0
                    continue
                text = "".join(node.text or "" for node in child.findall(".//w:t", ns))
                has_break = child.find(".//w:br", ns) is not None
                if not text.strip() and not has_break:
                    consecutive_empty += 1
                    max_consecutive_empty = max(max_consecutive_empty, consecutive_empty)
                else:
                    consecutive_empty = 0
            self.assertLessEqual(max_consecutive_empty, 1)

            body_children = list(body)
            last_content = body_children[-2]  # final child is w:sectPr
            self.assertIsNone(last_content.find(".//w:br[@w:type='page']", ns))
            self.assertTrue(
                "".join(node.text or "" for node in last_content.findall(".//w:t", ns)).strip()
            )
            for row in root.findall(".//w:tr", ns):
                self.assertIsNotNone(row.find("w:trPr/w:cantSplit", ns))

            doc = Document(path)
            for heading in ("Heading 1", "Heading 2", "Heading 3"):
                self.assertTrue(doc.styles[heading].paragraph_format.keep_with_next)

    def test_markdown_guide_matches_word_contract(self) -> None:
        text = (SETTLED / "本课型怎么填.md").read_text()
        for required in (
            "业务只提供三类信息",
            "中文页签大纲与页序",
            "能力来源解释",
            "同一课件只锁一个 style pack",
            "A / B / C 中性组合参考",
            "A · 7 页",
            "B · 6 页",
            "C · 5 页",
            "内容与中文页签确认 → 视觉与非商品插图绑定确认 → 正式商品图及授权确认",
            "我不填写内部 JSON 或页型 ID",
        ):
            self.assertIn(required, text)
        self.assertNotIn("## 推荐板块（可删可增）", text)
        self.assertNotIn("默认主路径", text)
        self.assertNotIn("金银花露", text)


if __name__ == "__main__":
    unittest.main()
