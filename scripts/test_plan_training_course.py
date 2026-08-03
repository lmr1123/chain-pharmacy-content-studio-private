import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


MODULE_PATH = Path(__file__).with_name("plan_training_course.py")
SPEC = importlib.util.spec_from_file_location("training_planner", MODULE_PATH)
PLANNER = importlib.util.module_from_spec(SPEC)
assert SPEC.loader
sys.modules[SPEC.name] = PLANNER
SPEC.loader.exec_module(PLANNER)


class TrainingPlannerTest(unittest.TestCase):
    def test_product_sections_choose_different_recipes_in_source_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "product.md"
            path.write_text(
                "# 测试商品培训\n"
                "## 需求教育\n已审核需求内容。\n"
                "## 核心功效\n已审核功效内容。\n"
                "## 适宜人群\n已审核人群内容。\n",
                encoding="utf-8",
            )
            manifest = PLANNER.build_manifest(
                path, "style-pack.reference-product-blue-v1"
            )
            self.assertEqual(
                [scene["intent"] for scene in manifest["scenes"]],
                ["need_education", "efficacy_evidence", "audience_sequence"],
            )
            self.assertEqual(manifest["unresolved_sections"], [])
            self.assertTrue(
                manifest["style_cohesion_policy"]["single_style_pack"]
            )

    def test_docx_ignores_instructional_paragraphs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "product.docx"
            document = Document()
            ignore = document.styles.add_style("Courseware Ignore", 1)
            paragraph = document.add_paragraph(style=ignore)
            paragraph.add_run("这段填写说明不进入脚本")
            document.add_heading("Word 商品培训", level=1)
            document.add_heading("产品特点", level=2)
            document.add_paragraph("已审核特点原文。")
            document.save(path)
            manifest = PLANNER.build_manifest(
                path, "style-pack.reference-product-blue-v1"
            )
            self.assertEqual(manifest["project_title"], "Word 商品培训")
            self.assertEqual(len(manifest["scenes"]), 1)
            self.assertEqual(
                manifest["scenes"][0]["intent"], "feature_evidence"
            )
            self.assertNotIn(
                "填写说明",
                manifest["scenes"][0]["slots"]["approved_text"],
            )

    def test_docx_maps_business_fields_and_keeps_unknown_content_visible(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "sleep.docx"
            document = Document()
            ignore = document.styles.add_style("Courseware Ignore", 1)
            document.add_paragraph("填写说明：替换示例文字。", style=ignore)
            document.add_heading("睡眠健康培训", level=1)
            document.add_heading("典型症状", level=2)
            document.add_paragraph("审核正文")
            document.add_paragraph("入睡困难、夜间易醒，白天容易疲倦。")
            document.add_paragraph("必须原样上屏的事实／数据／短文案")
            document.add_paragraph("入睡困难")
            document.add_paragraph("夜间易醒")
            document.add_paragraph("授权素材附件或文件名")
            document.add_paragraph("睡眠症状插画_已授权.png")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "业务补充备注"
            table.cell(0, 1).text = "数字仍待确认"
            document.save(path)

            manifest = PLANNER.build_manifest(
                path, "style-pack.reference-product-blue-v1"
            )
            slots = manifest["scenes"][0]["slots"]
            self.assertEqual(
                slots["approved_text"],
                "入睡困难、夜间易醒，白天容易疲倦。",
            )
            self.assertEqual(
                slots["on_screen_facts"], ["入睡困难", "夜间易醒"]
            )
            self.assertEqual(
                slots["authorized_assets"], ["睡眠症状插画_已授权.png"]
            )
            self.assertEqual(
                slots["unmapped_content"], ["业务补充备注：数字仍待确认"]
            )
            self.assertNotIn("替换示例文字", slots["approved_text"])

    def test_second_theme_changes_content_without_changing_style_pack(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            first = Path(temp_dir) / "first.md"
            second = Path(temp_dir) / "second.md"
            first.write_text(
                "# 辅酶课程\n## 核心功效\n能量生成。\n",
                encoding="utf-8",
            )
            second.write_text(
                "# 睡眠课程\n"
                "## 症状表现\n"
                "审核正文：入睡困难，夜间容易醒。\n"
                "必须上屏的事实/数据/短文案：入睡困难\n"
                "授权素材：睡眠症状插画.png\n",
                encoding="utf-8",
            )
            first_manifest = PLANNER.build_manifest(
                first, "style-pack.reference-product-blue-v1"
            )
            second_manifest = PLANNER.build_manifest(
                second, "style-pack.reference-product-blue-v1"
            )
            self.assertEqual(
                first_manifest["style_pack_id"],
                second_manifest["style_pack_id"],
            )
            self.assertNotEqual(
                first_manifest["scenes"][0]["scene_recipe_id"],
                second_manifest["scenes"][0]["scene_recipe_id"],
            )
            self.assertEqual(
                second_manifest["scenes"][0]["slots"]["on_screen_facts"],
                ["入睡困难"],
            )
            self.assertEqual(
                second_manifest["scenes"][0]["asset_matches"][0]["role_id"],
                "symptom_illustration",
            )

    def test_html_preview_contains_review_fields_and_gaps(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "product.md"
            path.write_text(
                "# 商品课程\n"
                "## 商品亮相\n"
                "审核正文：这是商品介绍。\n"
                "必须上屏事实/数据/短文案：商品名称\n",
                encoding="utf-8",
            )
            manifest = PLANNER.build_manifest(
                path, "style-pack.reference-product-blue-v1"
            )
            preview = PLANNER.render_storyboard_preview(manifest)
            self.assertIn("审核原文", preview)
            self.assertIn("屏幕短文案", preview)
            self.assertIn("推荐画面配方", preview)
            self.assertIn("素材缺口", preview)
            self.assertIn("authorized_product_packshot", preview)

    def test_markdown_bracket_labels_remain_compatible(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "structured.txt"
            path.write_text(
                "# 商品课程\n"
                "## 商品亮相\n"
                "【审核旁白原文】\n这是审核原文。\n"
                "【必须原样上屏的事实／短文案】\n- 商品名称\n"
                "【本章节授权素材】\n- 商品包装：pack.png\n",
                encoding="utf-8",
            )
            manifest = PLANNER.build_manifest(
                path, "style-pack.reference-product-blue-v1"
            )
            slots = manifest["scenes"][0]["slots"]
            self.assertEqual(slots["approved_text"], "这是审核原文。")
            self.assertEqual(slots["on_screen_facts"], ["商品名称"])
            self.assertEqual(
                slots["authorized_assets"], ["商品包装：pack.png"]
            )


if __name__ == "__main__":
    unittest.main()
